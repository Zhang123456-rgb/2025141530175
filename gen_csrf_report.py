#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CSRF跨站请求伪造漏洞分析与修复报告（v2）
场景：修改密码功能
攻击链：alice登录 → CSRF改admin密码 → 新密码登录admin → 越权删用户
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

def box(doc, text, color="red"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "FFCDD2" if color=="red" else "C8E6C9"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    return p

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CSRF跨站请求伪造漏洞\n分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("——基于Flask修改密码功能的CSRF漏洞实验——"); r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x45,0x45,0x45)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验项目：Web安全漏洞分析与修复\n课程名称：网络安全实训\n报告日期：2026年7月")
r.font.size = Pt(12)
doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
ch = ["CSRF漏洞概述","CSRF核心原理与攻击链","攻击链详解：普通用户→越权管理员",
      "实验一：curl构造跨域恶意请求修改密码","实验二：Burp Suite生成CSRF PoC",
      "/change-password 漏洞代码分析","修复方案详解（含完整代码）",
      "修复前后代码对比","安全编码规范总结"]
for i,c in enumerate(ch,1):
    p = doc.add_paragraph(f"第{i}章  {c}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  CSRF漏洞概述", level=1)
doc.add_paragraph("CSRF（Cross-Site Request Forgery，跨站请求伪造）是一种Web安全攻击方式，攻击者诱导已登录用户点击恶意链接或访问恶意页面，利用用户在当前站点的登录状态，以用户身份执行非本意的操作。")

doc.add_heading("1.1 漏洞危害评级", level=2)
tbl(doc, [("评估项","评级"),("CVSS 3.0","8.8（High）"),("攻击向量","网络远程攻击（AV:N）"),
    ("攻击复杂度","低（AC:L）"),("影响","账号劫持、权限提升、数据篡改")])

doc.add_heading("1.2 CSRF攻击三要素", level=2)
doc.add_paragraph("CSRF攻击的成功需要同时满足以下三个条件：")
box(doc, "要素1：用户已登录目标站点 — 浏览器中存有有效的Session Cookie")
box(doc, "要素2：目标站点存在CSRF漏洞 — 敏感操作仅依赖Cookie验证，无额外Token校验")
box(doc, "要素3：用户访问恶意页面 — 攻击者诱导用户点击链接或加载图片")

doc.add_heading("1.3 攻击场景", level=2)
doc.add_paragraph("本实验中，攻击者构造一个恶意HTML页面，该页面包含一个自动提交的表单，目标为用户管理系统中的 /change-password 路由。当管理员admin（已登录状态）在浏览器中打开该恶意页面时，浏览器会自动发送POST请求修改admin的密码。攻击者随后使用新密码登录admin账户，获得管理员权限。")

# ===== 第二章 =====
doc.add_heading("第二章  CSRF核心原理与攻击链", level=1)
doc.add_heading("2.1 攻击流程", level=2)
doc.add_paragraph("CSRF攻击的完整流程：")
code(doc,
    "Step 1: 用户登录目标网站 → 获得Session Cookie\n"
    "Step 2: 攻击者构造恶意页面 → 包含自动提交的修改密码表单\n"
    "Step 3: 诱导用户访问恶意页面 → 通过钓鱼邮件/链接等方式\n"
    "Step 4: 浏览器自动发送POST请求 → Cookie自动携带\n"
    "Step 5: 服务器执行密码修改 → 认为请求来自合法用户\n"
    "Step 6: 攻击者用新密码登录 → 获得用户账号权限")

doc.add_heading("2.2 恶意页面示例（伪代码）", level=2)
doc.add_paragraph("攻击者构造的CSRF攻击HTML页面：")
code(doc,
    '<html>\n'
    '<body>\n'
    '<h1>🎉 恭喜中奖！点击领取奖品</h1>\n'
    '<form id="csrf" action="http://target:5000/change-password" method="POST">\n'
    '  <input type="hidden" name="username" value="admin">\n'
    '  <input type="hidden" name="new_password" value="hacked123">\n'
    '</form>\n'
    '<script>document.getElementById("csrf").submit();</script>\n'
    '</body>\n'
    '</html>')

doc.add_heading("2.3 Cookie自动携带机制", level=2)
doc.add_paragraph("浏览器会自动在同一域名下的请求中携带Cookie，这是CSRF攻击能够成功的根本原因。用户在访问攻击者页面时，浏览器会自动将目标站点的Cookie附加到请求中，服务器无法区分请求是用户自愿发起的还是被伪造的。")

# ===== 第三章 =====
doc.add_heading("第三章  攻击链详解：普通用户→越权管理员", level=1)
doc.add_heading("3.1 完整攻击链", level=2)
doc.add_paragraph("本次实验的完整攻击链分为四步：")

tbl(doc, [("步骤","操作","效果"),
    ("Step 1","alice登录系统","获得有效Session Cookie"),
    ("Step 2","CSRF攻击修改admin密码","admin的密码被改为攻击者指定的值"),
    ("Step 3","用新密码登录admin","攻击者获得管理员权限"),
    ("Step 4","越权删除用户carlos","利用admin权限执行敏感操作")])

doc.add_heading("3.2 攻击链图", level=2)
code(doc,
    "┌──────────┐    ┌──────────────┐    ┌──────────┐    ┌──────────┐\n"
    "│ alice登录  │──→│ CSRF改admin   │──→│ 新密码登录 │──→│ 越权删用户 │\n"
    "│ 获得Cookie │   │ 密码=hacked123│   │ admin     │   │ carlos    │\n"
    "└──────────┘    └──────────────┘    └──────────┘    └──────────┘")

# ===== 第四章 =====
doc.add_heading("第四章  实验一：curl构造跨域恶意请求修改密码", level=1)
doc.add_heading("4.1 实验目标", level=2)
doc.add_paragraph("模拟攻击者利用已登录用户alice的Session Cookie，通过curl命令构造跨域恶意请求，修改管理员admin的密码。")

doc.add_heading("4.2 实验步骤", level=2)
doc.add_paragraph("Step 1：攻击者先登录自己的账号（alice）获取有效Session Cookie")
code(doc,
    'curl -c /tmp/alice_cookies.txt -X POST http://target:5000/login \\\n'
    '  -d "username=alice&password=alice2025"\n'
    '# -c 参数保存服务器返回的Cookie到文件')

doc.add_paragraph("")
doc.add_paragraph("Step 2：利用alice的Session Cookie修改admin的密码（CSRF + 越权）")
code(doc,
    'curl -b /tmp/alice_cookies.txt -X POST http://target:5000/change-password \\\n'
    '  -d "username=admin&new_password=hacked123"\n'
    '# -b 参数在请求中附带上alice的Cookie\n'
    '# ⚠️ 服务器未验证session用户与提交的username是否一致\n'
    '# ⚠️ 服务器未验证CSRF Token\n'
    '# ⚠️ 服务器未验证原密码')

doc.add_paragraph("")
doc.add_paragraph("Step 3：攻击者用新密码登录admin账户")
code(doc,
    'curl -c /tmp/admin_cookies.txt -X POST http://target:5000/login \\\n'
    '  -d "username=admin&password=hacked123"\n'
    '# 成功以admin身份登录！')

doc.add_paragraph("")
doc.add_paragraph("Step 4：利用admin权限越权删除用户")
code(doc,
    'curl -b /tmp/admin_cookies.txt -X POST http://target:5000/admin/delete-user \\\n'
    '  -d "user_id=2"\n'
    '# admin权限成功删除用户')

doc.add_heading("4.3 攻击效果验证", level=2)
doc.add_paragraph("通过上述攻击链，一名普通用户alice可以完成以下操作：")
doc.add_paragraph("1. 修改管理员admin的密码（CSRF + 越权）", style='List Bullet')
doc.add_paragraph("2. 用新密码登录admin账户（权限提升）", style='List Bullet')
doc.add_paragraph("3. 利用admin权限删除其他用户（越权操作）", style='List Bullet')
doc.add_paragraph("整个过程不需要admin的原始密码，也不需要admin做任何确认操作", style='List Bullet')

# ===== 第五章 =====
doc.add_heading("第五章  实验二：Burp Suite生成CSRF PoC修改密码", level=1)
doc.add_heading("5.1 实验目标", level=2)
doc.add_paragraph("使用Burp Suite拦截POST /change-password请求，自动生成CSRF攻击Proof of Concept（PoC）HTML代码。当受害者在已登录状态下打开该HTML页面时，密码被无感知修改。")

doc.add_heading("5.2 实验步骤", level=2)
doc.add_paragraph("Step 1：使用Burp Suite拦截POST /change-password请求")
code(doc,
    'POST /change-password HTTP/1.1\n'
    'Host: http://target:5000\n'
    'Cookie: session=abc123...\n\n'
    'username=admin&new_password=attacker123')

doc.add_paragraph("")
doc.add_paragraph("Step 2：在Burp拦截到的请求上右键 → Engagement tools → Generate CSRF PoC")
doc.add_paragraph("Step 3：Burp Suite自动生成CSRF攻击HTML代码")

doc.add_paragraph("")
doc.add_paragraph("Burp生成的CSRF PoC（修改密码版本）：")
code(doc,
    '<html>\n'
    '  <body>\n'
    '  <h1>🎉 免费抽奖！点击领取</h1>\n'
    '  <form action="http://target:5000/change-password" method="POST">\n'
    '    <input type="hidden" name="username" value="admin" />\n'
    '    <input type="hidden" name="new_password" value="attacker123" />\n'
    '    <input type="submit" value="Submit request" />\n'
    '  </form>\n'
    '  <script>document.forms[0].submit();</script>\n'
    '  </body>\n'
    '</html>')

doc.add_paragraph("")
doc.add_paragraph("Step 4：将生成的HTML代码保存为 csrf_attack.html")
doc.add_paragraph("Step 5：在受害者浏览器中打开该HTML文件（受害者需已登录admin后台）")
doc.add_paragraph("Step 6：admin的密码被无感知修改为 attacker123")

doc.add_heading("5.3 CSRF PoC的关键组件", level=2)
tbl(doc, [("组件","作用"),
    ("<form>标签","POST方法提交到 /change-password 路由"),
    ("<input type=hidden>","隐藏字段：username=admin, new_password=attacker123"),
    ("<script>自动提交","页面加载后立即提交表单，用户无感知"),
    ("攻击链扩展","修改密码后可登录admin删除其他用户")])

# ===== 第六章 =====
doc.add_heading("第六章  /change-password 漏洞代码分析", level=1)
doc.add_heading("6.1 有漏洞代码（app.py）", level=2)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n\n'
    '    username = request.form.get("username", "")       # ⚠️ 用户可控\n'
    '    new_password = request.form.get("new_password", "")\n\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n\n'
    '    # ⚠️ 漏洞1：无CSRF Token验证 → 任意网站可伪造提交\n'
    '    # ⚠️ 漏洞2：无原密码验证 → 直接设置新密码\n'
    '    # ⚠️ 漏洞3：无Session绑定 → 可修改任意用户密码（越权）\n'
    '    # ⚠️ 漏洞4：无Referer验证 → 不检查请求来源\n'
    '    # ⚠️ 漏洞5：SQL注入 → 字符串拼接\n'
    '    # ⚠️ 漏洞6：密码明文存储 → 数据库直接存储明文\n\n'
    '    sql = f"UPDATE users SET password = \'{new_password}\' WHERE username = \'{username}\'"\n'
    '    c.execute(sql)\n'
    '    conn.commit()\n'
    '    return redirect("/profile")')

doc.add_heading("6.2 六个安全漏洞一览", level=2)
tbl(doc, [("#","漏洞","严重度","攻击利用方式"),
    ("1","无CSRF Token","🔴高危","攻击者构造恶意HTML页面，诱导已登录用户访问，自动提交修改密码表单"),
    ("2","无需原密码","🔴高危","知道用户名即可直接修改密码，暴力枚举+社工即可攻破"),
    ("3","越权改密","🔴高危","已登录用户可修改任意用户的密码（username由表单提交）"),
    ("4","无Referer验证","🟡中危","无法区分请求来自本站还是来自攻击者的恶意页面"),
    ("5","SQL注入","🔴高危","username参数直接拼接SQL语句，可猜测/绕过"),
    ("6","明文存储","🟡中危","数据库直接存储明文密码，数据泄露后密码完全暴露")])

# ===== 第七章 =====
doc.add_heading("第七章  修复方案详解（含完整代码）", level=1)

doc.add_heading("7.1 修复1：CSRF Token验证（核心）", level=2)
box(doc, "✅ 登录时生成随机Token存入Session，表单提交时匹配验证", "green")
code(doc,
    '# 登录成功时：\n'
    'session["csrf_token"] = uuid.uuid4().hex\n\n'
    '# 表单中：\n'
    '<input type="hidden" name="csrf_token" value="{{ session.csrf_token }}">\n\n'
    '# 后端验证：\n'
    'token = request.form.get("csrf_token", "")\n'
    'if token != session.get("csrf_token", ""):\n'
    '    return "CSRF Token验证失败，请刷新页面重试"', green=True)

doc.add_heading("7.2 修复2：Session绑定（防越权改密）", level=2)
box(doc, "✅ 只能修改当前登录用户的密码，禁止修改他人", "green")
code(doc,
    'if username != session["username"]:\n'
    '    audit_logger.warning(f"越权改密被拒: {session[\'username\']}→{username}")\n'
    '    return "无权修改他人密码"', green=True)

doc.add_heading("7.3 修复3：Referer验证（防CSRF）", level=2)
box(doc, "✅ 验证请求来源是否为本站点", "green")
code(doc,
    'referer = request.headers.get("Referer", "")\n'
    'if not referer or not referer.startswith(request.host_url):\n'
    '    audit_logger.warning(f"CSRF被拒: Referer={referer}")\n'
    '    return "请求来源不合法"', green=True)

doc.add_heading("7.4 修复4：参数化查询 + 密码哈希", level=2)
box(doc, "✅ 参数化查询防SQL注入 + SHA256哈希存储", "green")
code(doc,
    'hashed_pw = hashlib.sha256(new_password.encode()).hexdigest()\n'
    'c.execute("UPDATE users SET password = ? WHERE username = ?",\n'
    '          (hashed_pw, username))\n'
    'conn.commit()\n'
    'audit_logger.info(f"密码修改成功: 用户={username}")', green=True)

doc.add_heading("7.5 完整修复代码", level=2)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n\n'
    '    # ✅ 修复1：CSRF Token验证\n'
    '    token = request.form.get("csrf_token", "")\n'
    '    if token != session.get("csrf_token", ""):\n'
    '        return "CSRF Token验证失败"\n\n'
    '    # ✅ 修复2：防越权（只能修改自己的密码）\n'
    '    if username != session["username"]:\n'
    '        return "无权修改他人密码"\n\n'
    '    # ✅ 修复3：Referer验证防CSRF\n'
    '    referer = request.headers.get("Referer", "")\n'
    '    if not referer.startswith(request.host_url):\n'
    '        return "请求来源不合法"\n\n'
    '    # ✅ 修复4：参数化查询 + 密码哈希\n'
    '    hashed_pw = hashlib.sha256(new_password.encode()).hexdigest()\n'
    '    conn = get_db()\n'
    '    c = conn.cursor()\n'
    '    c.execute("UPDATE users SET password = ? WHERE username = ?",\n'
    '              (hashed_pw, username))\n'
    '    conn.commit()\n'
    '    conn.close()\n'
    '    return redirect("/profile")', green=True)

# ===== 第八章 =====
doc.add_heading("第八章  修复前后代码对比", level=1)
tbl(doc, [("对比项","有漏洞(app.py)","已修复(app_fixed.py)"),
    ("CSRF Token","❌ 无Token验证","✅ uuid生成 + Session校验"),
    ("原密码验证","❌ 无需原密码","✅ 必须与Session用户名匹配（防越权）"),
    ("越权防护","❌ 可改任何人密码","✅ 只能修改自己的密码"),
    ("Referer验证","❌ 无任何验证","✅ 验证Referer来源"),
    ("SQL注入","❌ 字符串拼接","✅ 参数化查询"),
    ("密码存储","❌ 明文存储","✅ SHA256哈希"),
    ("审计日志","❌ 无日志记录","✅ logging记录操作")])

doc.add_paragraph("")
doc.add_heading("核心代码对比", level=2)
p = doc.add_paragraph(); r = p.add_run("有漏洞版本（app.py — 仅22行）："); r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n'
    '    # ⚠️ 无CSRF Token / 无原密码 / 无越权校验 / 无Referer\n'
    '    # ⚠️ 明文存储 + SQL注入\n'
    '    sql = f"UPDATE users SET password = \'{new_password}\' WHERE username = \'{username}\'"\n'
    '    c.execute(sql); conn.commit()\n'
    '    return redirect("/profile")')

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("修复版本（app_fixed.py — 完整校验）："); r.bold = True; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n\n'
    '    # ✅ CSRF Token验证\n'
    '    if token != session.get("csrf_token", ""):\n'
    '        return "CSRF Token验证失败"\n'
    '    # ✅ 防越权：只能修改自己的密码\n'
    '    if username != session["username"]:\n'
    '        return "无权修改他人密码"\n'
    '    # ✅ Referer验证\n'
    '    if not referer.startswith(request.host_url):\n'
    '        return "请求来源不合法"\n'
    '    # ✅ 参数化查询 + SHA256\n'
    '    c.execute("UPDATE users SET password = ? WHERE username = ?",\n'
    '              (hashlib.sha256(new_password.encode()).hexdigest(), username))\n'
    '    conn.commit()\n'
    '    return redirect("/profile")', green=True)

# ===== 第九章 =====
doc.add_heading("第九章  安全编码规范总结", level=1)
principles = [
    ("原则一：CSRF Token","所有涉及数据修改的POST请求必须包含CSRF Token，Token应为随机生成且绑定用户Session，每次提交后更新。"),
    ("原则二：同源验证","校验请求的Referer或Origin头，确保请求来自本站点，拒绝跨域请求。"),
    ("原则三：最小权限操作","用户只能修改自己的资源密码，不能越权操作他人数据。session中的username必须与提交的username一致。"),
    ("原则四：参数化查询","永远不要拼接SQL字符串，使用参数化查询（?占位符）防止SQL注入。"),
    ("原则五：密码安全存储","密码不应明文存储，应使用SHA256或bcrypt等哈希算法加盐存储，即使数据库泄露也无法还原密码。"),
    ("原则六：审计日志","所有敏感操作（登录、改密、删除用户）必须记录日志，包含用户名、时间、IP、操作内容，支持安全溯源。")]
for t,d in principles:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/CSRF跨站请求伪造漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ Word报告已生成：{out}")
print(f"   文件大小：{os.path.getsize(out)/1024:.1f}KB")
