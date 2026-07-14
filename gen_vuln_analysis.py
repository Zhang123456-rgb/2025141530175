#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
用户管理系统 — 全量漏洞分析与修复报告
覆盖：SQL注入、CSRF、越权、文件包含、路径穿越、XSS、业务逻辑
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("用户管理系统\n全量漏洞分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("——基于Flask的Web安全综合实验——"); r.font.size = Pt(14)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("报告日期：2026年7月"); r.font.size = Pt(12)
doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
chs = ["漏洞总体概览","SQL注入漏洞分析与修复","CSRF跨站请求伪造漏洞分析与修复",
       "越权漏洞（IDOR+垂直越权）分析与修复","文件包含与路径穿越漏洞分析与修复",
       "文件上传漏洞分析与修复","业务逻辑漏洞分析与修复","信息泄露漏洞分析与修复",
       "修复前后代码对比","安全编码规范总结"]
for i,c in enumerate(chs,1):
    p = doc.add_paragraph(f"第{i}章  {c}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章：漏洞总体概览 =====
doc.add_heading("第一章  漏洞总体概览", level=1)
doc.add_paragraph("本次安全实验基于Flask用户管理系统，系统包含登录、注册、搜索、文件上传、个人中心、充值、商城、管理面板、动态页面加载、密码修改等功能。在开发过程中故意引入了多种安全漏洞，用于教学演示和修复实践。")

doc.add_heading("漏洞汇总表", level=2)
tbl(doc, [("编号","漏洞类型","影响路由","严重度"),
    ("V-01","SQL注入（登录）","POST /login","🔴高危"),
    ("V-02","SQL注入（搜索）","GET /search","🔴高危"),
    ("V-03","SQL注入（注册）","POST /register","🔴高危"),
    ("V-04","SQL注入（个人中心）","GET /profile","🔴高危"),
    ("V-05","SQL注入（充值）","POST /recharge","🔴高危"),
    ("V-06","SQL注入（改密）","POST /change-password","🔴高危"),
    ("V-07","CSRF（无Token）","POST /change-password","🔴高危"),
    ("V-08","越权改密（无Session绑定）","POST /change-password","🔴高危"),
    ("V-09","IDOR越权查看资料","GET /profile?user_id=","🔴高危"),
    ("V-10","垂直越权访问管理面板","GET /admin","🔴高危"),
    ("V-11","垂直越权删除用户","POST /admin/delete-user","🔴高危"),
    ("V-12","文件包含/路径穿越","GET /page?name=","🔴高危"),
    ("V-13","文件上传无类型检查","POST /upload","🔴高危"),
    ("V-14","文件上传原始文件名","POST /upload","🔴高危"),
    ("V-15","价格篡改（客户端定价）","POST /cart","🔴高危"),
    ("V-16","负值充值","POST /recharge","🟡中危"),
    ("V-17","密码明文存储","所有涉及密码的路由","🟡中危"),
    ("V-18","HTML注释泄露账号","登录页/首页","🟡低危"),
    ("V-19","Debug模式开启","全局","🟡中危"),
    ("V-20","无审计日志","所有敏感操作","🟡低危")])

# ===== 逐一分析每个漏洞 =====
doc.add_page_break()
doc.add_heading("第二章  SQL注入漏洞分析与修复", level=1)
doc.add_paragraph("SQL注入是最经典的Web安全漏洞之一，攻击者通过向SQL查询中注入恶意代码，绕过认证、窃取数据、甚至执行系统命令。")

sqli_list = [
    ("V-01 登录SQL注入", "POST /login", "高危",
     'sql = f"SELECT * FROM users WHERE username=\'{username}\' AND password=\'{password}\'"',
     "admin' OR '1'='1' -- 可绕过密码登录任意账号",
     "c.execute('SELECT * FROM users WHERE username=? AND password=?', (username, password))"),
    ("V-02 搜索SQL注入", "GET /search", "高危",
     'sql = f"SELECT * FROM users WHERE username LIKE \'%{keyword}%\'"',
     "' OR 1=1 -- 可爆出所有用户信息",
     "c.execute('SELECT * FROM users WHERE username LIKE ?', ('%'+keyword+'%',))"),
    ("V-06 改密SQL注入", "POST /change-password", "高危",
     'sql = f"UPDATE users SET password = \'{new_password}\' WHERE username = \'{username}\'"',
     "username=admin' -- 可修改任意用户密码",
     "c.execute('UPDATE users SET password=? WHERE username=?', (new_password, username))"),
]

for title, route, severity, vuln_code, exploit, fix_code in sqli_list:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    r = p.add_run(f"[{severity}] {route}")
    r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
    doc.add_paragraph("漏洞代码："); doc.add_paragraph(vuln_code)
    doc.add_paragraph("攻击方式："); doc.add_paragraph(exploit)
    doc.add_paragraph("修复代码："); doc.add_paragraph(fix_code)

# CSRF
doc.add_page_break()
doc.add_heading("第三章  CSRF跨站请求伪造漏洞分析与修复", level=1)
doc.add_paragraph("CSRF攻击利用用户已登录状态，诱导用户访问恶意页面，在用户不知情的情况下执行非授权操作。")

doc.add_heading("V-07 无CSRF Token验证", level=2)
doc.add_paragraph("漏洞路由：POST /change-password")
doc.add_paragraph("漏洞代码：表单中无CSRF Token隐藏字段，后端无Token校验")
doc.add_paragraph("")
doc.add_paragraph("攻击链（4步）：")
doc.add_paragraph("Step1：alice登录获取Session Cookie")
doc.add_paragraph("Step2：CSRF攻击修改admin密码（利用alice的Cookie）")
doc.add_paragraph("Step3：用新密码登录admin获得管理员权限")
doc.add_paragraph("Step4：越权删除用户carlos")
doc.add_paragraph("")
doc.add_paragraph("curl攻击命令：")
doc.add_paragraph(
    "# alice登录\n"
    "curl -c /tmp/c.txt -X POST http://target/login -d 'username=alice&password=alice2025'\n\n"
    "# CSRF改admin密码\n"
    "curl -b /tmp/c.txt -X POST http://target/change-password -d 'username=admin&new_password=hacked'\n\n"
    "# 登录admin\n"
    "curl -c /tmp/admin.txt -X POST http://target/login -d 'username=admin&password=hacked'\n\n"
    "# 删除用户\n"
    "curl -b /tmp/admin.txt -X POST http://target/admin/delete-user -d 'user_id=2'")
doc.add_paragraph("")
doc.add_paragraph("修复方案（三重防护）：")
doc.add_paragraph("1. CSRF Token：uuid.uuid4().hex 生成，存入Session，表单提交时校验")
doc.add_paragraph("2. Session绑定：if username != session['username']: 禁止越权改密")
doc.add_paragraph("3. Referer验证：检查请求来源是否为本站点")

doc.add_heading("V-08 越权改密（无Session绑定）", level=2)
doc.add_paragraph("漏洞：/change-password从表单接收username参数，但不验证是否与当前登录用户一致。任意已登录用户可修改任意用户的密码。")
doc.add_paragraph("修复：if username != session['username']: return '无权修改他人密码'")

# 越权
doc.add_page_break()
doc.add_heading("第四章  越权漏洞（IDOR+垂直越权）分析与修复", level=1)

doc.add_heading("V-09 IDOR越权查看资料", level=2)
doc.add_paragraph("漏洞路由：GET /profile?user_id=N")
doc.add_paragraph("漏洞：从URL参数获取user_id查询用户资料，不验证当前登录用户是否有权查看。修改user_id即可查看任意用户信息。")
doc.add_paragraph("利用：/profile?user_id=1 → 查看管理员资料（含邮箱、手机、余额）")
doc.add_paragraph("修复：查询目标用户后，校验role==admin或目标username==session['username']")

doc.add_heading("V-10/V-11 垂直越权", level=2)
doc.add_paragraph("漏洞路由：GET /admin, POST /admin/delete-user")
doc.add_paragraph("漏洞：仅检查了登录状态，未检查用户角色是否为管理员（admin）。普通用户可直接访问管理面板并删除其他用户。")
doc.add_paragraph("修复：c.execute('SELECT role FROM users WHERE username=?', (username,)) → 检查role==admin")

# 文件包含
doc.add_page_break()
doc.add_heading("第五章  文件包含与路径穿越漏洞分析与修复", level=1)

doc.add_heading("V-12 文件包含/路径穿越", level=2)
doc.add_paragraph("漏洞路由：GET /page?name=")
doc.add_paragraph("漏洞代码：name参数直接拼接到文件路径中，未过滤../符号")
doc.add_paragraph("")
doc.add_paragraph("利用方式：")
doc.add_paragraph("/page?name=../app.py → 读取应用源码")
doc.add_paragraph("/page?name=../data/users.db → 下载数据库文件")
doc.add_paragraph("/page?name=../../../../../../etc/passwd → 读取系统文件")
doc.add_paragraph("")
doc.add_paragraph("伪协议知识（PHP环境参考）：")
doc.add_paragraph("http:// → 远程文件包含(RFI)、内网扫描")
doc.add_paragraph("data:// → 内联数据执行代码")
doc.add_paragraph("php://filter → Base64编码读取源码")
doc.add_paragraph("php://input → POST数据执行代码")
doc.add_paragraph("")
doc.add_paragraph("修复方案（三重防护）：")
doc.add_paragraph("1. 白名单：ALLOWED_PAGES = {'help','about','faq'}")
doc.add_paragraph("2. 过滤路径穿越：if '..' in name: return '页面不存在'")
doc.add_paragraph("3. 路径规范化：os.path.abspath() + startswith()校验")

# ===== 第九章：修复前后代码对比 =====
doc.add_page_break()
doc.add_heading("第九章  修复前后代码对比", level=1)

tbl(doc, [("功能","漏洞版本(app.py)","修复版本(app_fixed.py)"),
    ("登录","f-string拼接SQL","参数化查询?占位符"),
    ("搜索","f-string拼接SQL","参数化查询?占位符"),
    ("注册","f-string拼接SQL","参数化查询?占位符"),
    ("个人中心","IDOR+SQL注入","Session验证+参数化查询"),
    ("充值","负值充值+SQL注入","正数校验+参数化查询"),
    ("改密","无CSRF/无校验/明文","Token+Session+Referer+哈希"),
    ("上传","无检查/原始文件名","白名单+UUID+MIME检测+htaccess"),
    ("商城","客户端提交价格","服务端从数据库获取价格"),
    ("管理面板","无权限校验","role==admin校验"),
    ("删除用户","无权限校验","role==admin+审计日志"),
    ("页面加载","路径穿越无过滤","白名单+路径规范化")])

doc.add_paragraph("")
doc.add_paragraph("核心修复原则总结：")
doc.add_paragraph("1. 所有SQL查询使用参数化查询（?占位符），绝不拼接字符串", style='List Bullet')
doc.add_paragraph("2. 所有POST敏感操作添加CSRF Token校验", style='List Bullet')
doc.add_paragraph("3. 所有操作绑定当前Session用户，禁止越权", style='List Bullet')
doc.add_paragraph("4. 所有文件路径使用白名单+规范化校验", style='List Bullet')
doc.add_paragraph("5. 所有密码哈希存储（SHA256/bcrypt）", style='List Bullet')
doc.add_paragraph("6. 所有敏感操作记录审计日志", style='List Bullet')
doc.add_paragraph("7. 所有管理功能检查用户角色（role==admin）", style='List Bullet')
doc.add_paragraph("8. 所有上传文件后缀白名单+UUID重命名+目录权限控制", style='List Bullet')

# ===== 第十章：安全编码规范 =====
doc.add_heading("第十章  安全编码规范总结", level=1)
principles = [
    ("原则一：输入不可信","用户提交的任何数据都不可信，必须经过严格校验、过滤、转义。"),
    ("原则二：参数化查询","所有SQL操作使用参数化查询，绝不拼接字符串。"),
    ("原则三：最小权限","用户只能操作自己的资源；管理功能检查role；进程不以root运行。"),
    ("原则四：纵深防御","不依赖单一防御措施，多层防护：Token+Referer+Session+白名单。"),
    ("原则五：安全默认配置","框架安全特性默认开启(debug=False)，密钥不硬编码，敏感信息不泄露。"),
    ("原则六：审计日志","所有敏感操作记录日志，支持安全事件溯源。")]
for t,d in principles:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/用户管理系统全量漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ 报告已生成：{out}")
print(f"   大小：{os.path.getsize(out)/1024:.1f}KB")
