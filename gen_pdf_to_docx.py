#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从PDF提取内容生成Word文档"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_docx_from_text(txt_path, output_path, title):
    """从文本文件生成Word文档"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()

    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Detect headings: 一、二、三、目录等
        if line.startswith('目  录') or line.startswith('目录'):
            doc.add_heading(line, level=1)
        elif line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or \
             line.startswith('四、') or line.startswith('五、') or line.startswith('六、') or \
             line.startswith('七、') or line.startswith('八、') or line.startswith('九、'):
            doc.add_heading(line, level=2)
        elif any(line.startswith(f'{i}.') for i in range(1,15)):
            doc.add_heading(line, level=3)
        elif '漏洞' in line and len(line) < 50:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
        elif line.startswith('#') or line.startswith('sql') or line.startswith('keyword') or \
             line.startswith('  #') or line.startswith('"') or line.startswith("'") or \
             'f"' in line or 'execute' in line.lower() or 'sql' in line.lower()[:5]:
            # Code-like content
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            r = p.add_run(line)
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            p = doc.add_paragraph(line)

    doc.save(output_path)
    print(f"✅ {output_path}")
    return len(text)

# SQL注入报告
create_docx_from_text(
    '/root/flask-app-vuln/SQL注入漏洞分析与修复报告.txt',
    '/root/flask-app-vuln/static/SQL注入漏洞分析与修复报告.docx',
    'SQL注入漏洞分析与修复报告'
)

# Flask安全漏洞修复报告
create_docx_from_text(
    '/root/flask-app-vuln/Flask_安全漏洞修复报告.txt',
    '/root/flask-app-vuln/static/Flask_安全漏洞修复报告.docx',
    'Flask安全漏洞修复报告'
)

# Flask实验报告v2
create_docx_from_text(
    '/root/flask-app-vuln/Flask_安全漏洞分析与修复实验报告_v2.txt',
    '/root/flask-app-vuln/static/Flask_安全漏洞分析与修复实验报告_v2.docx',
    'Flask安全漏洞分析与修复实验报告'
)

print("\n全部转换完成！")
